import datetime
from io import BytesIO
from zipfile import ZipFile

from django.http import HttpResponse
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from Object.models import *


def rename_row_in_table(row, number):
    row.cells[0].text = number
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

names_in_docx = {'полеНазваниеобъекта':'name_object', \
                 'полеКодобъекта':'kod_object', \
                 'полеЗаявитель':'zayv', \
                 'полеНомерпроекта':'Nomer_proekt', \
                 'полеГод':str(datetime.datetime.today().year),\
                 'полеФИОсварщикп':'svarshik1.fio', \
                 'полеФИОсварщикс':'svarshik2.fio', \
                 'полеДиаметрп':'svarshik1.diametr',\
                 'полеДиаметрс':'svarshik2.diametr',\
                 'полеКолвоп':'svarshik1.kolvo',\
                 'полеКолвос':'svarshik2.kolvo'}

def form_dop_ITD(name):

    object = Object.objects.get(name_object=name)

    if object.rostehnadzor:

        doc = docx.Document('Перечень лиц, участвующих в строительстве.docx')

        style = doc.styles['Normal']

        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)

        style_under = doc.styles.add_style('Under_style', WD_STYLE_TYPE.PARAGRAPH)
        style_under.font.underline = True
        style_under.font.name = 'Times New Roman'
        style_under.font.size = Pt(10)

        for index, table in enumerate(doc.tables):
            for index, row in enumerate(table.rows):
                for index, cell in enumerate(row.cells):
                    if 'полеНазваниеобъекта' in cell.text:
                        if object.name_object:
                            cell.text = cell.text.replace('полеНазваниеобъекта', object.name_object)
                            cell.paragraphs[0].style = style_under
                        else:
                            cell.text = cell.text.replace('полеНазваниеобъекта', '')


        doc2 = docx.Document('Акт нивелировки.docx')

        for index, paragraph in enumerate(doc2.paragraphs):
            for index, run in enumerate(paragraph.runs):
                if 'полеДатаукладки' in run.text:
                    if object.Data_ukl:
                        run.text = run.text.replace('полеДатаукладки',
                            datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатаукладки', '')

                if 'полеНазваниеобъекта'in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта',
                            object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

                if 'полеПроектнаяорг'in run.text:
                    if object.proektnaya_org:
                        run.text = run.text.replace('полеПроектнаяорг',
                            object.proektnaya_org)
                    else:
                        run.text = run.text.replace('полеПроектнаяорг', '')

                if 'полеТехндолжность' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.post:
                                run.text = run.text.replace('полеТехндолжность',
                                    object.tehnadzor.person.post)
                            else:
                                run.text = run.text.replace('полеТехндолжность', '')
                        else:
                            run.text = run.text.replace('полеТехндолжность', '')
                    else:
                        run.text = run.text.replace('полеТехндолжность', '')


                if 'полеТехнФИО' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.fio:
                                run.text = run.text.replace('полеТехнФИО',
                                    object.tehnadzor.person.fio)
                            else:
                                run.text = run.text.replace('полеТехнФИО', '')
                        else:
                            run.text = run.text.replace('полеТехнФИО', '')
                    else:
                        run.text = run.text.replace('полеТехнФИО', '')


                if 'полеПК' in run.text:
                    if object.rostehnadzor:
                        run.text = run.text.replace('полеПК',
                            f'ПК{object.rostehnadzor.pk1}+{object.rostehnadzor.pk1_diam}-ПК{object.rostehnadzor.pk2}+{object.rostehnadzor.pk2_diam}')
                    else:
                        run.text = run.text.replace('полеПК', 'ПК_+_-ПК_+_')

        f = BytesIO()
        doc.save(f)
        f2 = BytesIO()
        doc2.save(f2)

        style = doc.styles['Normal']

        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        doc3 = docx.Document('Акты допом.docx')

        style_for_name_object = doc3.styles.add_style('style_for_name_object', WD_STYLE_TYPE.PARAGRAPH)
        style_for_name_object.font.bold = True
        style_for_name_object.font.italic = True
        style_for_name_object.font.name = 'Times New Roman'
        style_for_name_object.font.size = Pt(12)

        style_for_teh_11 = doc3.styles.add_style('style_for_teh_11', WD_STYLE_TYPE.PARAGRAPH)
        style_for_teh_11.font.name = 'Times New Roman'
        style_for_teh_11.font.size = Pt(11)

        for index, table in enumerate(doc3.tables):
            for index, row in enumerate(table.rows):
                for index, cell in enumerate(row.cells):
                    if 'полеНазваниеобъекта' in cell.text:
                        if object.name_object:
                            cell.text = cell.text.replace('полеНазваниеобъекта', object.name_object)
                        else:
                            cell.text = cell.text.replace('полеНазваниеобъекта', '')

                        cell.paragraphs[0].style = style_for_name_object

                    if 'полеДатаукладки ' in cell.text:
                        if object.Data_ukl:
                            cell.text = cell.text.replace('полеДатаукладки',
                                datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                        else:
                            cell.text = cell.text.replace('полеДатаукладки', '')

                        cell.paragraphs[0].alignment = paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if 'полеТехнаддолжность' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.post:
                                    cell.text = cell.text.replace('полеТехнаддолжность',
                                        object.tehnadzor.person.post)
                                    cell.paragraphs[0].style = style_for_teh_11
                                else:
                                    cell.text = cell.text.replace('полеТехнаддолжность', '')
                            else:
                                cell.text = cell.text.replace('полеТехнаддолжность', '')
                        else:
                            cell.text = cell.text.replace('полеТехнаддолжность', '')

                    if 'полеТехнадзорФИО' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.fio:
                                    cell.text = cell.text.replace('полеТехнадзорФИО',
                                        object.tehnadzor.person.fio)
                                    cell.paragraphs[0].style = style_for_teh_11
                                else:
                                    cell.text = cell.text.replace('полеТехнадзорФИО', '')
                            else:
                                cell.text = cell.text.replace('полеТехнадзорФИО', '')
                        else:
                            cell.text = cell.text.replace('полеТехнадзорФИО', '')

                    if 'полеТехнадзор2ФИО' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.fio:
                                    cell.text = cell.text.replace('полеТехнадзор2ФИО',
                                        object.tehnadzor.person.fio)
                                    cell.paragraphs[0].style = style_for_teh_11
                                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                else:
                                    cell.text = cell.text.replace('полеТехнадзор2ФИО', '')
                            else:
                                cell.text = cell.text.replace('полеТехнадзор2ФИО', '')
                        else:
                            cell.text = cell.text.replace('полеТехнадзор2ФИО', '')


                    if 'полеНомерпроекта' in cell.text:
                        if object.Nomer_proekt:
                            cell.text = cell.text.replace('полеНомерпроекта', object.Nomer_proekt)
                        else:
                            cell.text = cell.text.replace('полеНомерпроекта', '')

                    if 'полеПроектнаяорганизация' in cell.text:
                        if object.proektnaya_org:
                            cell.text = cell.text.replace('полеПроектнаяорганизация', object.proektnaya_org)
                        else:
                            cell.text = cell.text.replace('полеПроектнаяорганизация', '')

                    if 'полеТруба' in cell.text:
                        if object.rostehnadzor:
                            if object.rostehnadzor.truba:
                                cell.text = cell.text.replace('полеТруба',
                                    f'Труба {object.rostehnadzor.truba.diametr}x{object.rostehnadzor.truba.x}')
                            else:
                                cell.text = cell.text.replace('полеТруба', '')
                        else:
                            cell.text = cell.text.replace('полеТруба', '')

                    if 'полеДатаукладки1' in cell.text:
                        if object.Data_ukl:
                            cell.text = cell.text.replace('полеДатаукладки1',
                                datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                        else:
                            cell.text = cell.text.replace('полеДатаукладки1', '')

                    if 'полеДатаукладки2' in cell.text:
                        if object.Data_ukl:
                            cell.text = cell.text.replace('полеДатаукладки2',
                                datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                        else:
                            cell.text = cell.text.replace('полеДатаукладки2','')

        f3 = BytesIO()
        doc3.save(f3)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)

        doc4 = docx.Document('Акт входного контроля качества.docx')

        for index, paragraph in enumerate(doc4.paragraphs):
            for index, run in enumerate(paragraph.runs):

                if 'полеДатазамера' in run.text:
                    if object.Data_zamera2:
                        run.text = run.text.replace('полеДатазамера',
                            datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатазамера', '')

                if 'полеТехнадзордолжность' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.post:
                                run.text = run.text.replace('полеТехнадзордолжность',
                                    object.tehnadzor.person.post)
                            else:
                                run.text = run.text.replace('полеТехнадзордолжность', '')
                        else:
                            run.text = run.text.replace('полеТехнадзордолжность', '')
                    else:
                        run.text = run.text.replace('полеТехнадзордолжность', '')

                if 'полеТехнадзорФИО' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.fio:
                                run.text = run.text.replace('полеТехнадзорФИО',
                                    object.tehnadzor.person.fio)
                            else:
                                run.text = run.text.replace('полеТехнадзорФИО', '')
                        else:
                            run.text = run.text.replace('полеТехнадзорФИО', '')
                    else:
                        run.text = run.text.replace('полеТехнадзорФИО', '')

                if 'полеНомерпроекта' in run.text:
                    if object.Nomer_proekt:
                        run.text = run.text.replace('полеНомерпроекта', object.Nomer_proekt)
                    else:
                        run.text = run.text.replace('полеНомерпроекта', '')

                if 'полеПроектнаяорганизация' in run.text:
                    if object.proektnaya_org:
                        run.text = run.text.replace('полеПроектнаяорганизация', object.proektnaya_org)
                    else:
                        run.text = run.text.replace('полеПроектнаяорганизация', '')

                if 'полеНазваниеобъекта' in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта', object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

        f4 = BytesIO()
        doc4.save(f4)

    if object.prodavl:

        doc5 = docx.Document('Протокол бурения.docx')

        for index, paragraph in enumerate(doc5.paragraphs):
            for index, run in enumerate(paragraph.runs):

                if 'полеНазваниеобъекта' in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта', object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

                if 'полеДиамтр' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.diametr:
                                run.text = run.text.replace('полеДиамтр', object.prodavl.truba.diametr)
                            else:
                                run.text = run.text.replace('полеДиамтр', '')
                        else:
                            run.text = run.text.replace('полеДиамтр', '')
                    else:
                        run.text = run.text.replace('полеДиамтр', '')

                if 'полеДлнтр' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.dlina:
                                run.text = run.text.replace('полеДлнтр', object.prodavl.truba.dlina)
                            else:
                                run.text = run.text.replace('полеДлнтр', '')
                        else:
                            run.text = run.text.replace('полеДлнтр', '')
                    else:
                        run.text = run.text.replace('полеДлнтр', '')

                if 'полеПК1' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_1_diam:
                            run.text = run.text.replace('полеПК1', object.prodavl.pk0_1_diam)
                        else:
                            run.text = run.text.replace('полеПК1', '')
                    else:
                        run.text = run.text.replace('полеПК1', '')

                if 'полеПК2' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_2_diam:
                            run.text = run.text.replace('полеПК2', object.prodavl.pk0_2_diam)
                        else:
                            run.text = run.text.replace('полеПК2', '')
                    else:
                        run.text = run.text.replace('полеПК2', '')

                if 'полеДатаукл' in run.text:
                    if object.Data_ukl:
                        run.text = run.text.replace('полеДатаукл',
                            datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатаукл','')

        f5 = BytesIO()
        doc5.save(f5)

        doc6 = docx.Document('Акт приёмки.docx')
        style_for_te = doc6.styles.add_style('style_for_teh_11', WD_STYLE_TYPE.PARAGRAPH)
        style_for_te.font.name = 'Times New Roman'
        style_for_te.font.size = Pt(11)


        for index, paragraph in enumerate(doc6.paragraphs):
            for index, run in enumerate(paragraph.runs):

                if 'полеДатаукладки' in run.text:
                    if object.Data_ukl:
                        run.text = run.text.replace('полеДатаукладки',
                            datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатаукладки','')

                if 'полеНазваниеобъекта' in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта', object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

                if 'полеТехнФИО' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.fio:
                                run.text = run.text.replace('полеТехнФИО',
                                    object.tehnadzor.person.fio)
                            else:
                                run.text = run.text.replace('полеТехнФИО', '')
                        else:
                            run.text = run.text.replace('полеТехнФИО', '')
                    else:
                        run.text = run.text.replace('полеТехнФИО', '')

                if 'полеТрх' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.x:
                                run.text = run.text.replace('полеТрх', object.prodavl.truba.x)
                            else:
                                run.text = run.text.replace('полеТрх', '')
                        else:
                            run.text = run.text.replace('полеТрх', '')
                    else:
                        run.text = run.text.replace('полеТрх', '')


                if 'полеТрдиам' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.diametr:
                                run.text = run.text.replace('полеТрдиам', object.prodavl.truba.diametr)
                            else:
                                run.text = run.text.replace('полеТрдиам', '')
                        else:
                            run.text = run.text.replace('полеТрдиам', '')
                    else:
                        run.text = run.text.replace('полеТрдиам', '')

                if 'полеТрдл' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.dlina:
                                run.text = run.text.replace('полеТрдл', object.prodavl.truba.dlina)
                            else:
                                run.text = run.text.replace('полеТрдл', '')
                        else:
                            run.text = run.text.replace('полеТрдл', '')
                    else:
                        run.text = run.text.replace('полеТрдл', '')

                if 'полеПК1' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_1_diam:
                            run.text = run.text.replace('полеПК1', object.prodavl.pk0_1_diam)
                        else:
                            run.text = run.text.replace('полеПК1', '')
                    else:
                        run.text = run.text.replace('полеПК1', '')

                if 'полеПК2' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_2_diam:
                            run.text = run.text.replace('полеПК2', object.prodavl.pk0_2_diam)
                        else:
                            run.text = run.text.replace('полеПК2', '')
                    else:
                        run.text = run.text.replace('полеПК2', '')

                if 'полеПроектнаяорг' in run.text:
                    if object.proektnaya_org:
                        run.text = run.text.replace('полеПроектнаяорг', object.proektnaya_org)
                    else:
                        run.text = run.text.replace('полеПроектнаяорг', '')

                if 'полеНомерпроекта' in run.text:
                    if object.Nomer_proekt:
                        run.text = run.text.replace('полеНомерпроекта', object.Nomer_proekt)
                    else:
                        run.text = run.text.replace('полеНомерпроекта', '')

        for index, table in enumerate(doc6.tables):
            for index, row in enumerate(table.rows):
                for index, cell in enumerate(row.cells):
                    if 'полеТехнФИО' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.fio:
                                    cell.text = cell.text.replace('полеТехнФИО',
                                        object.tehnadzor.person.fio)
                                    cell.paragraphs[0].style = style_for_te
                                else:
                                    cell.text = cell.text.replace('полеТехнФИО', '')
                            else:
                                cell.text = cell.text.replace('полеТехнФИО', '')
                        else:
                            cell.text = cell.text.replace('полеТехнФИО', '')

                    if 'полеТехнаддолж' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.post:
                                    cell.text = cell.text.replace('полеТехнаддолж',
                                        object.tehnadzor.person.post)
                                    cell.paragraphs[0].style = style_for_te
                                else:
                                    cell.text = cell.text.replace('полеТехнаддолж', '')
                            else:
                                cell.text = cell.text.replace('полеТехнаддолж', '')
                        else:
                            cell.text = cell.text.replace('полеТехнаддолж', '')

        f6 = BytesIO()
        doc6.save(f6)

    response = HttpResponse(content_type='application/zip')
    zf = ZipFile(response, 'w')

    zf.writestr('Перечень лиц, участвующих в строительстве.docx', f.getvalue())
    zf.writestr('Акт нивелировки.docx', f2.getvalue())
    zf.writestr('Акты допом.docx', f3.getvalue())
    zf.writestr('Акт входного контроля качества.docx', f4.getvalue())
    zf.writestr('Протокол бурения.docx', f5.getvalue())
    zf.writestr('Акт приёмки.docx', f6.getvalue())

    response['Content-Disposition'] = f'attachment; filename=dop_ITD.zip'
    return response


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
