import datetime
from io import BytesIO
from zipfile import ZipFile

from django.http import HttpResponse
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from Object.models import *

def rename_row_in_table(row, number):
    row.cells[0].text = number
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

names_in_docx = {'полеНазваниеобъекта':'name_object', \
                 'полеКодобъекта':'kod_object', \
                 'полеЗаявитель':'zayv', \
                 'полеНомерпроекта':'Nomer_proekt', \
                 'полеГод':str(datetime.datetime.today().year),\
                 'полеФИОсварщикп':'svarshik1.fio', \
                 'полеФИОсварщикс':'svarshik2.fio', \
                 'полеДиаметрп':'svarshik1.diametr',\
                 'полеДиаметрс':'svarshik2.diametr',\
                 'полеКолвоп':'svarshik1.kolvo',\
                 'полеКолвос':'svarshik2.kolvo'}

def form_ITD(name):
    object = Object.objects.get(name_object=name)

    doc = docx.Document('ITD_template.docx')

    flag_prodavlivanie = False

    if not object.prodavlivanie:
        remove_row(doc.tables[0], doc.tables[0].rows[14])
        for row_index in range(14,18):
            rename_row_in_table(doc.tables[0].rows[row_index], str(row_index))
        flag_prodavlivanie = True

    flag_izolir_soed = False

    if object.gazoprovod_nadzem_stal:
        if object.gazoprovod_nadzem_stal.izolir_soed:
            if not object.gazoprovod_nadzem_stal.izolir_soed.diametr:
                remove_row(doc.tables[0], doc.tables[0].rows[9])
                flag_izolir_soed = True
                if not flag_prodavlivanie:
                    for row_index in range(9,18):
                        rename_row_in_table(doc.tables[0].rows[row_index], str(row_index))
                else:
                    for row_index in range(9,17):
                        rename_row_in_table(doc.tables[0].rows[row_index], str(row_index))

    if not object.gazoprovod:
        remove_row(doc.tables[0], doc.tables[0].rows[3])
        if not flag_prodavlivanie:
            if not flag_izolir_soed:
                for row_index in range(3,18):
                    rename_row_in_table(doc.tables[0].rows[row_index], str(row_index))
            else:
                for row_index in range(3,17):
                    rename_row_in_table(doc.tables[0].rows[row_index], str(row_index))
        else:
            if not flag_izolir_soed:
                for row_index in range(3,17):
                    rename_row_in_table(doc.tables[0].rows[row_index], str(row_index))
            else:
                for row_index in range(3,16):
                    rename_row_in_table(doc.tables[0].rows[row_index], str(row_index))

    for index, row in enumerate(doc.tables[0].rows):
        if index == 0:
            continue
        row.cells[0].text = str(index)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    style = doc.styles['Normal']

    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    style_bold = doc.styles.add_style('Bold_style', WD_STYLE_TYPE.PARAGRAPH)
    style_bold.font.bold = True
    style_bold.font.size = Pt(11)

    style_11 = doc.styles.add_style('Normal_11', WD_STYLE_TYPE.PARAGRAPH)
    style_11.font.size = Pt(11)

    new_paragraph = doc.add_paragraph()
    new_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:

            """
            ГАЗОПРОВОДЫ
            """

            if 'Списокгазопроводов' in run.text:
                paragraph.text = ''

                paragraph.add_run("{}\n\n".format(object.name_gazoprovod)).bold = True

                if object.gazoprovod_podzem_stal:
                    if object.gazoprovod_podzem_stal.truba:
                        paragraph.add_run("{}\n".format(object.name_podzem_stal_gazoprovod)).bold = True
                        for truba in object.gazoprovod_podzem_stal.truba.all():
                            if truba.diametr:
                                if truba.prim:
                                    paragraph.add_run("- Труба Ø{}x{} - {} м {}\n".format(
                                        truba.diametr,
                                        truba.x,
                                        truba.dlina,
                                        f'({truba.prim})'))
                                else:
                                    paragraph.add_run("- Труба Ø{}x{} - {} м \n".format(
                                    truba.diametr,
                                    truba.x,
                                    truba.dlina))

                    paragraph.add_run("  Установлено:\n")

                    if object.gazoprovod_podzem_stal.neraz_soed:
                        if object.gazoprovod_podzem_stal.neraz_soed.PE:
                            paragraph.add_run("    - Неразъемное соединение ПЭØ{}/СТØ{} - {} шт.\n".format(
                                object.gazoprovod_podzem_stal.neraz_soed.PE,
                                object.gazoprovod_podzem_stal.neraz_soed.ST,
                                object.gazoprovod_podzem_stal.neraz_soed.kolvo))

                    if object.gazoprovod_podzem_stal.kontrolnaya_trubka:
                        paragraph.add_run("    - Контрольная трубка - {} шт.\n".format(
                            object.gazoprovod_podzem_stal.kontrolnaya_trubka))

                    if object.gazoprovod_podzem_stal.otvod_90:
                        paragraph.add_run("    - Отвод ст. 90 гр. - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.otvod_90))

                    if object.gazoprovod_podzem_stal.opoznavat_znak:
                        paragraph.add_run("    - Опознавательный знак - {} шт.\n".format(
                            object.gazoprovod_podzem_stal.opoznavat_znak))

                if object.gazoprovod_podzem_poliet:
                    if object.gazoprovod_podzem_poliet.truba:
                        paragraph.add_run("\n{}\n".format(object.name_podzem_poliet_gazoprovod)).bold = True
                        try:
                            for truba in object.gazoprovod_podzem_poliet.truba.all():
                                if truba.diametr:
                                    paragraph.add_run("- Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                                        truba.diametr, truba.x, truba.dlina))
                        except:
                            for truba in object.gazoprovod_podzem_poliet.truba.all():
                                if truba.diametr:
                                    paragraph.add_run("- Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                                        truba.diametr, truba.x, truba.dlina))

                    paragraph.add_run("  Установлено:\n")

                    if object.gazoprovod_podzem_poliet.mufta:
                        for mufta in object.gazoprovod_podzem_poliet.mufta.all():
                            if mufta.diametr:
                                paragraph.add_run("    - Муфта электросварная ПЭ100 SDR11 Ø{}х{} – {} шт.\n".format(
                                    mufta.diametr, mufta.x, mufta.dlina))

                    if object.gazoprovod_podzem_poliet.otvod:
                        if object.gazoprovod_podzem_poliet.otvod.diametr:
                            paragraph.add_run("    - Отвод ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.otvod.diametr,
                                object.gazoprovod_podzem_poliet.otvod.kolvo))

                    if object.gazoprovod_podzem_poliet.troinik:
                        if object.gazoprovod_podzem_poliet.troinik.diametr1:
                            paragraph.add_run("    - Тройник электросварной ПЭ100 SDR11 Ø{}xØ{}xØ{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.troinik.diametr1,
                                object.gazoprovod_podzem_poliet.troinik.diametr2,
                                object.gazoprovod_podzem_poliet.troinik.diametr3,
                                object.gazoprovod_podzem_poliet.troinik.kolvo))

                    if object.gazoprovod_podzem_poliet.zaglushka:
                        if object.gazoprovod_podzem_poliet.zaglushka.diametr:
                            paragraph.add_run("    - Заглушка электросварная ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.zaglushka.diametr,
                                object.gazoprovod_podzem_poliet.zaglushka.kolvo))

                    if object.gazoprovod_podzem_poliet.lenta:
                        paragraph.add_run("    - Лента сигнальная “Осторожно! ГАЗ!”  – {} м\n".format(
                            object.gazoprovod_podzem_poliet.lenta))

                    if object.gazoprovod_podzem_poliet.kran:
                        if object.gazoprovod_podzem_poliet.kran.diametr:
                            paragraph.add_run("    - Кран шаровой для подземной установки ПЭ 100 ГАЗ SDR11 Ø{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.kran.diametr,
                                object.gazoprovod_podzem_poliet.kran.kolvo))

                    if object.gazoprovod_podzem_poliet.znak:
                        paragraph.add_run("    - Опознавательный знак – {} шт.\n".format(
                            object.gazoprovod_podzem_poliet.znak))

                    if object.gazoprovod_podzem_poliet.sedelka:
                        if object.gazoprovod_podzem_poliet.sedelka.diametr1:
                            paragraph.add_run("    - Седелка электросварная Ø{}xØ{}xØ{} – {} шт.\n".format(
                            object.gazoprovod_podzem_poliet.sedelka.diametr1,
                            object.gazoprovod_podzem_poliet.sedelka.diametr2,
                            object.gazoprovod_podzem_poliet.sedelka.diametr3,
                            object.gazoprovod_podzem_poliet.sedelka.kolvo))

                if object.gazoprovod_nadzem_stal:
                    paragraph.add_run("\n{}\n".format(object.name_nadzem_stal_gazoprovod)).bold = True

                    if object.gazoprovod_nadzem_stal.truba:
                        for truba in object.gazoprovod_nadzem_stal.truba.all():
                            if truba.diametr:
                                paragraph.add_run("- Труба стальная Ø{}х{} – {} м\n".format(
                                    truba.diametr, truba.x, truba.dlina))

                    paragraph.add_run("  Установлено:\n")

                    if object.gazoprovod_nadzem_stal.izolir_soed:
                        if object.gazoprovod_nadzem_stal.izolir_soed.diametr:
                            paragraph.add_run("    - Изолирующие соединения Ø{} – {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.izolir_soed.diametr,
                                object.gazoprovod_nadzem_stal.izolir_soed.kolvo))

                    if object.gazoprovod_nadzem_stal.kran_stal:
                        if object.gazoprovod_nadzem_stal.kran_stal.diametr:
                            paragraph.add_run("    - Кран стальной Ø{} – {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.kran_stal.diametr,
                                object.gazoprovod_nadzem_stal.kran_stal.kolvo))

                    if object.gazoprovod_nadzem_stal.otvod:
                        if object.gazoprovod_nadzem_stal.otvod.diametr:
                            paragraph.add_run("    - Отвод Ø{} – {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.otvod.diametr,
                                object.gazoprovod_nadzem_stal.otvod.kolvo))

                    if object.gazoprovod_nadzem_stal.cokolnyi_vvod:
                        if object.gazoprovod_nadzem_stal.cokolnyi_vvod.PE:
                            paragraph.add_run("    - Цокольный ввод ПЭØ{}/СТØ{} - {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.cokolnyi_vvod.PE,
                                object.gazoprovod_nadzem_stal.cokolnyi_vvod.ST,
                                object.gazoprovod_nadzem_stal.cokolnyi_vvod.kolvo))

                    if object.gazoprovod_nadzem_stal.kreplenie:
                        paragraph.add_run("    - Крепление газопровода к кирпичной стене – {} шт.\n".format(
                            object.gazoprovod_nadzem_stal.kreplenie))

                    if object.gazoprovod_nadzem_stal.stoika:
                        if object.gazoprovod_nadzem_stal.stoika.diametr:
                            paragraph.add_run("    - Стойка под газопровод Ø{} – {} м - {} шт.\n\n".format(
                                object.gazoprovod_nadzem_stal.stoika.diametr,
                                object.gazoprovod_nadzem_stal.stoika.dlina,
                                object.gazoprovod_nadzem_stal.stoika.kolvo))

                paragraph.add_run("\nИтого протяженность газопровода – {} м\n".format(
                    object.itogo)).bold = True

                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.style = style_11

            """
            СЕРТИФИКАТЫ
            """

            if 'Списоксертификатов' in run.text:

                count = 1
                paragraph.text = ''
                if object.certificates:
                    for cert in object.certificates.all():
                        paragraph.add_run('{}. {}\n'.format(count, cert.name))
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        count += 1

            """
            Страница 1
            """

            for name in names_in_docx:
                if name != 'полеГод':
                    try:
                        if len(name.split('.')) > 0:
                            itog_field = object
                            for index in range(0,len(names_in_docx[name].split('.'))):
                                splitted_field = names_in_docx[name].split('.')[index]
                                itog_field = getattr(itog_field,splitted_field)
                            if itog_field:
                                run.text = run.text.replace(name,itog_field)
                            else:
                                run.text = run.text.replace(name, '')
                        else:
                            if getattr(object,names_in_docx[name]):
                                run.text = run.text.replace(name,getattr(object,names_in_docx[name]))
                            else:
                                run.text = run.text.replace(name, '')
                    except:
                        pass
                else:
                    run.text = run.text.replace(name,names_in_docx[name])

    """
    Страница 3
    """
    try:
        if object.kod_object:
            doc.tables[1].rows[0].cells[1].text = doc.tables[1].rows[0].cells[1].text.replace(
                'полеНомерпроекта', object.kod_object)
        else:
            doc.tables[1].rows[0].cells[1].text = doc.tables[1].rows[0].cells[1].text.replace(
                'полеНомерпроекта', '')

        if object.name_object:
            doc.tables[1].rows[3].cells[0].text = doc.tables[1].rows[3].cells[0].text.replace(
                'полеНазваниеобъекта', object.name_object)
        else:
            doc.tables[1].rows[3].cells[0].text = doc.tables[1].rows[3].cells[0].text.replace(
                'полеНазваниеобъекта', '')

        if object.kod_object:
            doc.tables[1].rows[3].cells[0].text = doc.tables[1].rows[3].cells[0].text.replace(
                'полеКодобъекта', object.kod_object)
        else:
            doc.tables[1].rows[3].cells[0].text = doc.tables[1].rows[3].cells[0].text.replace(
                'полеКодобъекта', '')

        for paragraph in doc.tables[1].rows[3].cells[0].paragraphs:
            paragraph.style = style_bold
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass

    if object.tehnadzor:
        if object.tehnadzor.person.post:
            doc.tables[4].rows[7].cells[4].text = doc.tables[4].rows[7].cells[4].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[4].rows[18].cells[4].text = doc.tables[4].rows[18].cells[4].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[5].rows[12].cells[0].text = doc.tables[5].rows[12].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[5].rows[23].cells[0].text = doc.tables[5].rows[23].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[6].rows[11].cells[0].text = doc.tables[6].rows[11].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[8].rows[1].cells[0].text = doc.tables[8].rows[1].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[12].rows[7].cells[0].text = doc.tables[12].rows[7].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[15].rows[0].cells[0].text = doc.tables[15].rows[0].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[16].rows[11].cells[9].text = doc.tables[16].rows[11].cells[9].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[16].rows[28].cells[7].text = doc.tables[16].rows[28].cells[7].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[17].rows[11].cells[9].text = doc.tables[17].rows[11].cells[9].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[17].rows[29].cells[7].text = doc.tables[17].rows[29].cells[7].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[19].rows[5].cells[0].text = doc.tables[19].rows[5].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[21].rows[6].cells[0].text = doc.tables[21].rows[6].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[22].rows[10].cells[8].text = doc.tables[22].rows[10].cells[8].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[22].rows[21].cells[18].text = doc.tables[22].rows[21].cells[18].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[24].rows[5].cells[0].text = doc.tables[24].rows[5].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
            doc.tables[25].rows[14].cells[0].text = doc.tables[25].rows[14].cells[0].text.replace(
                'полеТехнадзордолжность', object.tehnadzor.person.post)
        else:
            doc.tables[4].rows[7].cells[4].text = doc.tables[4].rows[7].cells[4].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[4].rows[18].cells[4].text = doc.tables[4].rows[18].cells[4].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[5].rows[12].cells[0].text = doc.tables[5].rows[12].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[5].rows[23].cells[0].text = doc.tables[5].rows[23].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[6].rows[11].cells[0].text = doc.tables[6].rows[11].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[8].rows[1].cells[0].text = doc.tables[8].rows[1].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[12].rows[7].cells[0].text = doc.tables[12].rows[7].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[15].rows[0].cells[0].text = doc.tables[15].rows[0].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[16].rows[11].cells[9].text = doc.tables[16].rows[11].cells[9].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[16].rows[28].cells[7].text = doc.tables[16].rows[28].cells[7].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[17].rows[11].cells[9].text = doc.tables[17].rows[11].cells[9].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[17].rows[29].cells[7].text = doc.tables[17].rows[29].cells[7].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[19].rows[5].cells[0].text = doc.tables[19].rows[5].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[21].rows[6].cells[0].text = doc.tables[21].rows[6].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[22].rows[10].cells[8].text = doc.tables[22].rows[10].cells[8].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[22].rows[21].cells[18].text = doc.tables[22].rows[21].cells[18].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[24].rows[5].cells[0].text = doc.tables[24].rows[5].cells[0].text.replace(
                'полеТехнадзордолжность', '')
            doc.tables[25].rows[14].cells[0].text = doc.tables[25].rows[14].cells[0].text.replace(
                'полеТехнадзордолжность', '')

        if object.tehnadzor.person.fio:
            paragraph2 = doc.tables[4].rows[7].cells[4].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph3 = doc.tables[4].rows[7].cells[4].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph4 = doc.tables[4].rows[18].cells[4].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph5 = doc.tables[4].rows[18].cells[4].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph6 = doc.tables[5].rows[12].cells[0].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph7 = doc.tables[5].rows[12].cells[0].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph8= doc.tables[5].rows[23].cells[0].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph8.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph9 = doc.tables[5].rows[23].cells[0].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph9.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph10 = doc.tables[8].rows[1].cells[0].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph10.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph11 = doc.tables[8].rows[1].cells[0].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph11.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph12 = doc.tables[12].rows[7].cells[0].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph13 = doc.tables[12].rows[7].cells[0].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph13.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph14 = doc.tables[15].rows[0].cells[0].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph14.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph15 = doc.tables[15].rows[0].cells[0].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph15.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph16 = doc.tables[16].rows[28].cells[7].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph16.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph17 = doc.tables[16].rows[28].cells[7].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph17.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph18 = doc.tables[17].rows[29].cells[7].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph18.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph19 = doc.tables[17].rows[29].cells[7].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph19.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph20 = doc.tables[21].rows[6].cells[0].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph20.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21 = doc.tables[21].rows[6].cells[0].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph22 = doc.tables[22].rows[21].cells[18].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph23 = doc.tables[22].rows[21].cells[18].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph23.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph24 = doc.tables[25].rows[14].cells[0].add_paragraph(
                '_________________________       {}'.format(
                object.tehnadzor.person.fio))
            paragraph24.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph25 = doc.tables[25].rows[14].cells[0].add_paragraph(
                '(должность производителя работ, фамилия, имя, отчество)')
            paragraph25.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.tables[6].rows[11].cells[0].text = doc.tables[6].rows[11].cells[0].text.replace(
                'полеТехнадзорФИО', object.tehnadzor.person.fio)
            doc.tables[16].rows[11].cells[9].text = doc.tables[16].rows[11].cells[9].text.replace(
                'полеТехнадзорФИО', object.tehnadzor.person.fio)
            doc.tables[17].rows[11].cells[9].text = doc.tables[17].rows[11].cells[9].text.replace(
                'полеТехнадзорФИО', object.tehnadzor.person.fio)
            doc.tables[19].rows[5].cells[0].text = doc.tables[19].rows[5].cells[0].text.replace(
                'полеТехнадзорФИО', object.tehnadzor.person.fio)
            doc.tables[22].rows[10].cells[8].text = doc.tables[22].rows[10].cells[8].text.replace(
                'полеТехнадзорФИО', object.tehnadzor.person.fio)
            doc.tables[24].rows[5].cells[0].text = doc.tables[24].rows[5].cells[0].text.replace(
                'полеТехнадзорФИО', object.tehnadzor.person.fio)
        else:
            doc.tables[6].rows[11].cells[0].text = doc.tables[6].rows[11].cells[0].text.replace(
                'полеТехнадзорФИО', '')
            doc.tables[16].rows[11].cells[9].text = doc.tables[16].rows[11].cells[9].text.replace(
                'полеТехнадзорФИО', '')
            doc.tables[17].rows[11].cells[9].text = doc.tables[17].rows[11].cells[9].text.replace(
                'полеТехнадзорФИО', '')
            doc.tables[19].rows[5].cells[0].text = doc.tables[19].rows[5].cells[0].text.replace(
                'полеТехнадзорФИО', '')
            doc.tables[22].rows[10].cells[8].text = doc.tables[22].rows[10].cells[8].text.replace(
                'полеТехнадзорФИО', '')
            doc.tables[24].rows[5].cells[0].text = doc.tables[24].rows[5].cells[0].text.replace(
                'полеТехнадзорФИО', '')

    if object.Data_zamera1:
        doc.tables[4].rows[21].cells[5].text = doc.tables[4].rows[21].cells[5].text.replace(
            'полеДатазамера1', datetime.datetime.strftime(object.Data_zamera1,"%d.%m.%Y"))
        doc.tables[5].rows[3].cells[0].text = doc.tables[5].rows[3].cells[0].text.replace(
            'полеЧисло1', datetime.datetime.strftime(object.Data_zamera1,"%d"))
        doc.tables[5].rows[3].cells[1].text = doc.tables[5].rows[3].cells[1].text.replace(
            'полеМесяц1', datetime.datetime.strftime(object.Data_zamera1,"%m"))
        doc.tables[5].rows[3].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[5].rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    else:
        doc.tables[4].rows[21].cells[5].text = doc.tables[4].rows[21].cells[5].text.replace(
            'полеДатазамера1', '')
        doc.tables[5].rows[3].cells[0].text = doc.tables[5].rows[3].cells[0].text.replace(
            'полеЧисло1', '')
        doc.tables[5].rows[3].cells[1].text = doc.tables[5].rows[3].cells[1].text.replace(
            'полеМесяц1', '')


    if object.Data_zamera2:
        doc.tables[4].rows[23].cells[5].text = doc.tables[4].rows[23].cells[5].text.replace(
            'полеДатазамера2', datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
        doc.tables[5].rows[4].cells[0].text = doc.tables[5].rows[4].cells[0].text.replace(
            'полеЧисло2', datetime.datetime.strftime(object.Data_zamera2,"%d"))
        doc.tables[5].rows[4].cells[1].text = doc.tables[5].rows[4].cells[1].text.replace(
            'полеМесяц2', datetime.datetime.strftime(object.Data_zamera2,"%m"))
        doc.tables[5].rows[4].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[5].rows[4].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.tables[4].rows[23].cells[5].text = doc.tables[4].rows[23].cells[5].text.replace(
            'полеДатазамера2', '')
        doc.tables[5].rows[4].cells[0].text = doc.tables[5].rows[4].cells[0].text.replace(
            'полеЧисло2', '')
        doc.tables[5].rows[4].cells[1].text = doc.tables[5].rows[4].cells[1].text.replace(
            'полеМесяц2', '')

    if object.prover_davl:
        doc.tables[5].rows[3].cells[3].text = doc.tables[5].rows[3].cells[3].text.replace(
            'полеДавление1', object.prover_davl)
        doc.tables[5].rows[3].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[5].rows[4].cells[5].text = doc.tables[5].rows[4].cells[5].text.replace(
            'полеДавление2', object.prover_davl)
        doc.tables[5].rows[4].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[5].rows[6].cells[3].text = doc.tables[5].rows[6].cells[3].text.replace(
            'полеДавление3', object.prover_davl)
        doc.tables[5].rows[6].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.tables[5].rows[3].cells[3].text = doc.tables[5].rows[3].cells[3].text.replace(
            'полеДавление1', '')
        doc.tables[5].rows[4].cells[5].text = doc.tables[5].rows[4].cells[5].text.replace(
            'полеДавление2', '')
        doc.tables[5].rows[6].cells[3].text = doc.tables[5].rows[6].cells[3].text.replace(
            'полеДавление3', '')

    if object.proektnaya_org:
        doc.tables[5].rows[16].cells[0].text = doc.tables[5].rows[16].cells[0].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
    else:
        doc.tables[5].rows[16].cells[0].text = doc.tables[5].rows[16].cells[0].text.replace(
            'полеПроектнаяорганизация', '')

    if object.Data_sost_project:
        doc.tables[5].rows[16].cells[0].text = doc.tables[5].rows[16].cells[0].text.replace(
            'полеДатасоставленияобъекта', datetime.datetime.strftime(
            object.Data_sost_project,"%d.%m.%Y"))
        doc.tables[5].rows[16].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.tables[5].rows[16].cells[0].text = doc.tables[5].rows[16].cells[0].text.replace(
            'полеДатасоставленияобъекта', '')

    if object.Data_razbiv:
        doc.tables[5].rows[20].cells[4].text = doc.tables[5].rows[20].cells[4].text.replace(
            'полеДатаразбивки', datetime.datetime.strftime(object.Data_razbiv,"%d.%m.%Y"))
    else:
        doc.tables[5].rows[20].cells[4].text = doc.tables[5].rows[20].cells[4].text.replace(
            'полеДатаразбивки', '')

    if object.Data_zamera2:
        doc.tables[5].rows[21].cells[4].text = doc.tables[5].rows[21].cells[4].text.replace(
            'полеДатазамера2', datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
    else:
        doc.tables[5].rows[21].cells[4].text = doc.tables[5].rows[21].cells[4].text.replace(
            'полеДатазамера2', '')

    if object.Data_produv:
        doc.tables[6].rows[1].cells[6].text = doc.tables[6].rows[1].cells[6].text.replace(
            'полеДатапродувки', datetime.datetime.strftime(object.Data_produv,"%d.%m.%Y"))
        doc.tables[6].rows[1].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[6].rows[1].cells[6].text = doc.tables[6].rows[1].cells[6].text.replace(
            'полеДатапродувки', '')

    if object.name_object:
        doc.tables[6].rows[3].cells[0].text = doc.tables[6].rows[3].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
    else:
        doc.tables[6].rows[3].cells[0].text = doc.tables[6].rows[3].cells[0].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[6].rows[3].cells[0].text = doc.tables[6].rows[3].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[6].rows[3].cells[0].text = doc.tables[6].rows[3].cells[0].text.replace(
            'полеКодобъекта', '')

    if object.proektnaya_org:
        doc.tables[7].rows[1].cells[0].text = doc.tables[7].rows[1].cells[0].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
    else:
        doc.tables[7].rows[1].cells[0].text = doc.tables[7].rows[1].cells[0].text.replace(
            'полеПроектнаяорганизация', '')

    if object.Nomer_proekt:
        doc.tables[7].rows[1].cells[4].text = doc.tables[7].rows[1].cells[4].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
    else:
        doc.tables[7].rows[1].cells[4].text = doc.tables[7].rows[1].cells[4].text.replace(
            'полеНомерпроекта', '')

    if object.Data_zamera2:
        doc.tables[9].rows[0].cells[5].text = doc.tables[9].rows[0].cells[5].text.replace(
            'полеДатазамера2', datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
        doc.tables[9].rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[9].rows[0].cells[5].text = doc.tables[9].rows[0].cells[5].text.replace(
            'полеДатазамера2', '')

    if object.Nomer_proekt:
        doc.tables[10].rows[0].cells[3].text = doc.tables[10].rows[0].cells[3].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
    else:
        doc.tables[10].rows[0].cells[3].text = doc.tables[10].rows[0].cells[3].text.replace(
            'полеНомерпроекта', '')

    if object.name_object:
        doc.tables[10].rows[2].cells[0].text = doc.tables[10].rows[2].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
    else:
        doc.tables[10].rows[2].cells[0].text = doc.tables[10].rows[2].cells[0].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[10].rows[2].cells[0].text = doc.tables[10].rows[2].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[10].rows[2].cells[0].text = doc.tables[10].rows[2].cells[0].text.replace(
            'полеКодобъекта', '')

    if object.Data_razbiv:
        doc.tables[13].rows[0].cells[5].text = doc.tables[13].rows[0].cells[5].text.replace(
            'полеДатаразбивки1', datetime.datetime.strftime(object.Data_razbiv,"%d.%m.%Y"))
        doc.tables[13].rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[13].rows[0].cells[5].text = doc.tables[13].rows[0].cells[5].text.replace(
            'полеДатаразбивки1', '')

    if object.Data_razbiv:
        doc.tables[14].rows[0].cells[0].text = doc.tables[14].rows[0].cells[0].text.replace(
            'полеДатаразбивки2', datetime.datetime.strftime(object.Data_razbiv,"%d.%m.%Y"))
    else:
        doc.tables[14].rows[0].cells[0].text = doc.tables[14].rows[0].cells[0].text.replace(
            'полеДатаразбивки2', '')

    if object.Nomer_proekt:
        doc.tables[14].rows[1].cells[1].text = doc.tables[14].rows[1].cells[1].text.replace(
            'полеНомеробъекта', object.Nomer_proekt)
    else:
        doc.tables[14].rows[1].cells[1].text = doc.tables[14].rows[1].cells[1].text.replace(
            'полеНомеробъекта', '')

    if object.proektnaya_org:
        doc.tables[14].rows[1].cells[4].text = doc.tables[14].rows[1].cells[4].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
    else:
        doc.tables[14].rows[1].cells[4].text = doc.tables[14].rows[1].cells[4].text.replace(
            'полеПроектнаяорганизация', '')

    if object.name_object:
        doc.tables[14].rows[4].cells[0].text = doc.tables[14].rows[4].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
        doc.tables[14].rows[4].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.tables[14].rows[4].cells[0].text = doc.tables[14].rows[4].cells[0].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[14].rows[4].cells[0].text = doc.tables[14].rows[4].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[14].rows[4].cells[0].text = doc.tables[14].rows[4].cells[0].text.replace(
            'полеКодобъекта', '')

    if object.district:

        if object.district.name == "Автозаводский район":
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[15].cells[0].text.replace(
            'полеЗаказчикразбив',
            'Главный инженер Автозаводского производственного управления ПАО «Газпром газораспределение Нижний Новгород»')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________       {}'.format(
            'Камбаратов А.В'))
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif object.district.name == "Сормовский район":
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[15].cells[0].text.replace(
            'полеЗаказчикразбив', 'Начальник СПУ ПАО «Газпром газораспределение Нижний Новгород»')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________       {}'.format(
            'Метелев М.Л.'))
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif object.district.name == "Нижегородский район":
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[16].cells[0].text.replace(
            'полеЗаказчикразбив',
            'Главный инженер Нагорного производственного управления ПАО «Газпром газораспределение Нижний Новгород»')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________       {}'.format(
            'Котихин А.Н.'))
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        else:
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[15].cells[0].text.replace(
            'полеЗаказчикразбив', '')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________                             ')
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if object.name_object:
        doc.tables[16].rows[1].cells[0].text = doc.tables[16].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
    else:
        doc.tables[16].rows[1].cells[0].text = doc.tables[16].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[16].rows[1].cells[0].text = doc.tables[16].rows[1].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[16].rows[1].cells[0].text = doc.tables[16].rows[1].cells[0].text.replace(
            'полеКодобъекта', '')

    if object.Data_ukl:
        doc.tables[16].rows[3].cells[8].text = doc.tables[16].rows[3].cells[8].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[16].rows[3].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[16].rows[3].cells[8].text = doc.tables[16].rows[3].cells[8].text.replace(
            'полеДатаукладки', '')

    doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
        'Работы выполнены по проектно-сметной документации',
        '2.  Работы выполнены по проектно-сметной документации')

    if object.Nomer_proekt:
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
    else:
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеНомерпроекта', '')

    if object.proektnaya_org:
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
    else:
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеПроектнаяорганизация', '')

    if object.Data_sost_project:
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(
            object.Data_sost_project,"%d.%m.%Y"))
    else:
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеДатасоставленияпроекта', '')

    doc.tables[16].rows[18].cells[7].text = doc.tables[16].rows[18].cells[7].text.replace(
        'При выполнении работ применены', '3.  При выполнении работ применены')

    if object.dop_dann_sharovyi_kran:
        kran_mufta_string = ''
        for kran_mufta in object.dop_dann_sharovyi_kran.all():
            if kran_mufta.kran.diametr:
                substr = 'ПЭ100 SDR11 кран Ø{}; Муфта ПЭ 100 ГАЗ SDR11 Ø{}; '.format(
                kran_mufta.kran.diametr, kran_mufta.mufta.diametr)
                kran_mufta_string += substr
        doc.tables[16].rows[19].cells[7].text = doc.tables[16].rows[19].cells[7].text.replace(
            'полеКранмуфта,', kran_mufta_string)

    doc.tables[16].rows[20].cells[7].text = doc.tables[16].rows[20].cells[7].text.replace(
        'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

    if object.Data_ukl:
        doc.tables[16].rows[22].cells[7].text = doc.tables[16].rows[22].cells[7].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[16].rows[22].cells[7].text = doc.tables[16].rows[22].cells[7].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[16].rows[22].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.tables[16].rows[22].cells[7].text = doc.tables[16].rows[22].cells[7].text.replace(
            'полеДатаукладки1', '')
        doc.tables[16].rows[22].cells[7].text = doc.tables[16].rows[22].cells[7].text.replace(
            'полеДатаукладки2', '')

    if object.name_object:
        doc.tables[17].rows[1].cells[0].text = doc.tables[17].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
    else:
        doc.tables[17].rows[1].cells[0].text = doc.tables[17].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[17].rows[1].cells[0].text = doc.tables[17].rows[1].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[17].rows[1].cells[0].text = doc.tables[17].rows[1].cells[0].text.replace(
            'полеКодобъекта', '')

    if object.Data_ukl:
        doc.tables[17].rows[3].cells[8].text = doc.tables[17].rows[3].cells[8].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[17].rows[3].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[17].rows[3].cells[8].text = doc.tables[17].rows[3].cells[8].text.replace(
            'полеДатаукладки', '')

    doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
        'Работы выполнены по проектно-сметной документации',
        '2.  Работы выполнены по проектно-сметной документации')

    if object.Nomer_proekt:
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
    else:
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеНомерпроекта', '')

    if object.proektnaya_org:
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
    else:
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеПроектнаяорганизация', '')

    if object.Data_sost_project:
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(
            object.Data_sost_project,"%d.%m.%Y"))
    else:
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеДатасоставленияпроекта', '')

    doc.tables[17].rows[18].cells[7].text = doc.tables[17].rows[18].cells[7].text.replace(
        'При выполнении работ применены', '3.  При выполнении работ применены')

    doc.tables[17].rows[21].cells[7].text = doc.tables[17].rows[21].cells[7].text.replace(
        'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

    if object.Data_ukl:
        doc.tables[17].rows[23].cells[7].text = doc.tables[17].rows[23].cells[7].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[17].rows[23].cells[7].text = doc.tables[17].rows[23].cells[7].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[17].rows[23].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.tables[17].rows[23].cells[7].text = doc.tables[17].rows[23].cells[7].text.replace(
            'полеДатаукладки1', '')
        doc.tables[17].rows[23].cells[7].text = doc.tables[17].rows[23].cells[7].text.replace(
            'полеДатаукладки2', '')

    if object.name_object:
        doc.tables[18].rows[1].cells[0].text = doc.tables[18].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
    else:
        doc.tables[18].rows[1].cells[0].text = doc.tables[18].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[18].rows[1].cells[0].text = doc.tables[18].rows[1].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[18].rows[1].cells[0].text = doc.tables[18].rows[1].cells[0].text.replace(
            'полеКодобъекта', '')

    if object.Data_ukl:
        doc.tables[18].rows[3].cells[1].text = doc.tables[18].rows[3].cells[1].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[18].rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[18].rows[3].cells[1].text = doc.tables[18].rows[3].cells[1].text.replace(
            'полеДатаукладки', '')

    if object.zashitnyi_futlyar:
        if object.zashitnyi_futlyar.diametr:
            doc.tables[20].rows[3].cells[0].text = doc.tables[20].rows[3].cells[0].text.replace(
                'полеЗащитныйфутляр', 'Ø{}х{} L={} м'.format(object.zashitnyi_futlyar.diametr,
                object.zashitnyi_futlyar.x, object.zashitnyi_futlyar.dlina))

    doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
        'Работы выполнены по проектно-сметной документации',
        '2.  Работы выполнены по проектно-сметной документации')

    if object.Nomer_proekt:
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
    else:
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеНомерпроекта', '')

    if object.proektnaya_org:
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
    else:
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеПроектнаяорганизация', '')

    if object.Data_sost_project:
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(
            object.Data_sost_project,"%d.%m.%Y"))
    else:
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеДатасоставленияпроекта', '')

    doc.tables[20].rows[5].cells[0].text = doc.tables[20].rows[5].cells[0].text.replace(
        'При выполнении работ применены', '3.  При выполнении работ применены')

    if object.dop_dann_gazopr_v_zashit:
        trubi = ''
        for kran_mufta in object.dop_dann_gazopr_v_zashit.all():
            if kran_mufta.truba1.diametr or kran_mufta.truba2.diametr:
                substr = 'Труба ПЭ100 ГАЗ SDR11 Ø{}х{}; труба ст.Ø{}х{}; '.format(
                        kran_mufta.truba1.diametr, kran_mufta.truba1.x,
                        kran_mufta.truba2.diametr, kran_mufta.truba2.x)
                trubi += substr
        if trubi:
            trubi_endpoint = trubi[:-2] + '.'
        else:
            trubi_endpoint = ''
        doc.tables[20].rows[6].cells[0].text = doc.tables[20].rows[6].cells[0].text.replace(
            'полеТрубазащитныйфутляр', trubi_endpoint)

    doc.tables[20].rows[7].cells[0].text = doc.tables[20].rows[7].cells[0].text.replace(
        'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

    if object.Data_ukl:
        doc.tables[21].rows[0].cells[0].text = doc.tables[21].rows[0].cells[0].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[21].rows[0].cells[0].text = doc.tables[21].rows[0].cells[0].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[21].rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.tables[21].rows[0].cells[0].text = doc.tables[21].rows[0].cells[0].text.replace(
            'полеДатаукладки1', '')
        doc.tables[21].rows[0].cells[0].text = doc.tables[21].rows[0].cells[0].text.replace(
            'полеДатаукладки2', '')

    if object.name_object:
        doc.tables[22].rows[0].cells[18].text = doc.tables[22].rows[0].cells[18].text.replace(
            'полеНазваниеобъекта', object.name_object)
    else:
        doc.tables[22].rows[0].cells[18].text = doc.tables[22].rows[0].cells[18].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[22].rows[0].cells[18].text = doc.tables[22].rows[0].cells[18].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[22].rows[0].cells[18].text = doc.tables[22].rows[0].cells[18].text.replace(
            'полеКодобъекта', '')

    if object.Data_ukl:
        doc.tables[22].rows[3].cells[0].text = doc.tables[22].rows[3].cells[0].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[22].rows[3].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[22].rows[3].cells[0].text = doc.tables[22].rows[3].cells[0].text.replace(
            'полеДатаукладки', '')

    if object.futlyar_na_vyhode:
        if object.futlyar_na_vyhode.diametr:
            doc.tables[22].rows[14].cells[3].text = doc.tables[22].rows[14].cells[3].text.replace(
                'полеФутлярнавыходе', 'Ø{}х{} L={}м – {} шт'.format(object.futlyar_na_vyhode.diametr,
                object.futlyar_na_vyhode.x, object.futlyar_na_vyhode.dlina,
                object.futlyar_na_vyhode.kolvo))

    doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
        'Работы выполнены по проектно-сметной документации',
        '2.  Работы выполнены по проектно-сметной документации')

    if object.Nomer_proekt:
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
    else:
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'полеНомерпроекта', '')

    if object.proektnaya_org:
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
    else:
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'полеПроектнаяорганизация', '')

    if object.Data_sost_project:
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(
            object.Data_sost_project,"%d.%m.%Y"))
    else:
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'полеДатасоставленияпроекта', '')

    doc.tables[22].rows[15].cells[12].text = doc.tables[22].rows[15].cells[12].text.replace(
        'При выполнении работ применены', '3.  При выполнении работ применены')

    if object.dop_dann_futlyar_na_vyhode:
        trubi = ''
        for truba in object.dop_dann_futlyar_na_vyhode.all():
            if truba.diametr:
                substr = 'Труба Ø{}х{}; '.format(truba.diametr, truba.x)
                trubi += substr
        doc.tables[22].rows[16].cells[11].text = doc.tables[22].rows[16].cells[11].text.replace(
            'полеТрубафутлярнавыходе', trubi)

    doc.tables[22].rows[17].cells[10].text = doc.tables[22].rows[17].cells[10].text.replace(
        'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

    if object.Data_ukl:
        doc.tables[22].rows[18].cells[9].text = doc.tables[22].rows[18].cells[9].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[22].rows[18].cells[9].text = doc.tables[22].rows[18].cells[9].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[22].rows[18].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.tables[22].rows[18].cells[9].text = doc.tables[22].rows[18].cells[9].text.replace(
            'полеДатаукладки1', '')
        doc.tables[22].rows[18].cells[9].text = doc.tables[22].rows[18].cells[9].text.replace(
            'полеДатаукладки2', '')

    if object.name_object:
        doc.tables[23].rows[1].cells[0].text = doc.tables[23].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
    else:
        doc.tables[23].rows[1].cells[0].text = doc.tables[23].rows[1].cells[0].text.replace(
            'полеНазваниеобъекта', '')

    if object.kod_object:
        doc.tables[23].rows[1].cells[0].text = doc.tables[23].rows[1].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
    else:
        doc.tables[23].rows[1].cells[0].text = doc.tables[23].rows[1].cells[0].text.replace(
            'полеКодобъекта', '')

    if object.Data_ukl:
        doc.tables[23].rows[3].cells[1].text = doc.tables[23].rows[3].cells[1].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[23].rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.tables[23].rows[3].cells[1].text = doc.tables[23].rows[3].cells[1].text.replace(
            'полеДатаукладки', '')

    if object.opora:
        if object.opora.diametr:
            doc.tables[25].rows[3].cells[0].text = doc.tables[25].rows[3].cells[0].text.replace(
                'полеОпораподгазопровод',
                'монтаж опоры Ø{}х{} L={} м над уровнем земли под надземный газопровод низкого давления в количестве {} шт'.format(
                object.opora.diametr, object.opora.x, object.opora.dlina, object.opora.kolvo))

    doc.tables[25].rows[4].cells[0].text = doc.tables[25].rows[4].cells[0].text.replace(
        'Работы выполнены по проектно-сметной документации',
        '2.  Работы выполнены по проектно-сметной документации')
    doc.tables[25].rows[5].cells[0].text = doc.tables[25].rows[5].cells[0].text.replace(
        'При выполнении работ применены', '3.  При выполнении работ применены')

    if object.dop_dann_ob_ustanovke_opor:
        trubi = ''
        for truba in object.dop_dann_ob_ustanovke_opor.all():
            if truba.diametr:
                substr = 'Труба Ø{}х{}; '.format(truba.diametr, truba.x)
                trubi += substr
        doc.tables[25].rows[6].cells[0].text = doc.tables[25].rows[6].cells[0].text.replace(
            'полеТрубаопораподгаз', trubi)

    doc.tables[25].rows[7].cells[0].text = doc.tables[25].rows[7].cells[0].text.replace(
        'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

    if object.Data_ukl:
        doc.tables[25].rows[8].cells[0].text = doc.tables[25].rows[8].cells[0].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[25].rows[8].cells[0].text = doc.tables[25].rows[8].cells[0].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[25].rows[8].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.tables[25].rows[8].cells[0].text = doc.tables[25].rows[8].cells[0].text.replace(
            'полеДатаукладки1', '')
        doc.tables[25].rows[8].cells[0].text = doc.tables[25].rows[8].cells[0].text.replace(
            'полеДатаукладки2', '')

    flag_delete_svarshik2 = False
    if object.svarshik2:
        if object.svarshik2.fio:
            doc.tables[4].rows[4].cells[5].text = doc.tables[4].rows[4].cells[5].text.replace(
                'полеФИОсварщикс', object.svarshik2.fio)
            doc.tables[4].rows[5].cells[1].text = doc.tables[4].rows[5].cells[1].text.replace(
                'полеДиаметрс', object.svarshik2.diametr)
            doc.tables[4].rows[5].cells[2].text = doc.tables[4].rows[5].cells[2].text.replace(
                'полеКолвос', object.svarshik2.kolvo)
        else:
            flag_delete_svarshik2 = True
            doc.tables[4].rows[4].cells[5].text = doc.tables[4].rows[4].cells[5].text.replace(
                'полеФИОсварщикс', '')
            doc.tables[4].rows[5].cells[1].text = doc.tables[4].rows[5].cells[1].text.replace(
                'полеДиаметрс', '')
            doc.tables[4].rows[5].cells[2].text = doc.tables[4].rows[5].cells[2].text.replace(
                'полеКолвос', '')

        if object.svarshik2.date_svarki:
            doc.tables[4].rows[5].cells[3].text = doc.tables[4].rows[5].cells[3].text.replace(
                'полеДатас', datetime.datetime.strftime(object.svarshik2.date_svarki,"%d.%m.%Y"))
        else:
            doc.tables[4].rows[5].cells[3].text = doc.tables[4].rows[5].cells[3].text.replace(
                'полеДатас', '')
    else:
        flag_delete_svarshik2 = True
        doc.tables[4].rows[4].cells[5].text = doc.tables[4].rows[4].cells[5].text.replace(
            'полеФИОсварщикс', '')
        doc.tables[4].rows[5].cells[1].text = doc.tables[4].rows[5].cells[1].text.replace(
            'полеДиаметрс', '')
        doc.tables[4].rows[5].cells[2].text = doc.tables[4].rows[5].cells[2].text.replace(
            'полеКолвос', '')

    flag_delete_svarshik1 = False
    if object.svarshik1:
        if object.svarshik1.fio:
            doc.tables[4].rows[3].cells[6].text = doc.tables[4].rows[3].cells[6].text.replace(
                'полеФИОсварщикп', object.svarshik1.fio)
            doc.tables[4].rows[4].cells[2].text = doc.tables[4].rows[4].cells[2].text.replace(
                'полеДиаметрп', object.svarshik1.diametr)
            doc.tables[4].rows[4].cells[3].text = doc.tables[4].rows[4].cells[3].text.replace(
                'полеКолвоп', object.svarshik1.kolvo)
        else:
            flag_delete_svarshik1 = True
            doc.tables[4].rows[3].cells[6].text = doc.tables[4].rows[3].cells[6].text.replace(
                'полеФИОсварщикп', '')
            doc.tables[4].rows[4].cells[2].text = doc.tables[4].rows[4].cells[2].text.replace(
                'полеДиаметрп', '')
            doc.tables[4].rows[4].cells[3].text = doc.tables[4].rows[4].cells[3].text.replace(
                'полеКолвоп', '')

        if object.svarshik1.date_svarki:
            doc.tables[4].rows[4].cells[4].text = doc.tables[4].rows[4].cells[4].text.replace(
                'полеДатап', datetime.datetime.strftime(object.svarshik1.date_svarki,"%d.%m.%Y"))
        else:
            doc.tables[4].rows[4].cells[4].text = doc.tables[4].rows[4].cells[4].text.replace(
                'полеДатап', '')
    else:
        flag_delete_svarshik1 = True
        doc.tables[4].rows[3].cells[6].text = doc.tables[4].rows[3].cells[6].text.replace(
            'полеФИОсварщикп', '')
        doc.tables[4].rows[4].cells[2].text = doc.tables[4].rows[4].cells[2].text.replace(
            'полеДиаметрп', '')
        doc.tables[4].rows[4].cells[3].text = doc.tables[4].rows[4].cells[3].text.replace(
            'полеКолвоп', '')

    if object.svarshik2:
        if object.svarshik2.fio:
            if object.prover_davl == '0.3':
                doc.tables[11].rows[2].cells[0].text = 'Низкого давления (сталь)'
            elif object.prover_davl == '0.6':
                doc.tables[11].rows[2].cells[0].text = 'Среднего давления (сталь)'
            else:
                doc.tables[11].rows[2].cells[0].text = '{} (сталь)'.format(
                    object.davlenie_name.capitalize())

            doc.tables[11].rows[2].cells[1].text = "{} МПа".format(object.prover_davl)
            if object.Data_zamera1:
                doc.tables[11].rows[2].cells[2].text = "{} - {}".format(
                    datetime.datetime.strftime(object.Data_zamera1,"%d.%m.%Y"),
                    datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
            doc.tables[11].rows[2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.tables[11].rows[2].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            remove_row(doc.tables[11],doc.tables[11].rows[2])
    else:
        remove_row(doc.tables[11],doc.tables[11].rows[2])

    if object.svarshik1:
        if object.svarshik1.fio:
            if object.prover_davl == '0.3':
                doc.tables[11].rows[1].cells[0].text = 'Низкого давления (полиэтилен)'
            elif object.prover_davl == '0.6':
                doc.tables[11].rows[1].cells[0].text = 'Среднего давления (полиэтилен)'
            else:
                doc.tables[11].rows[1].cells[0].text = '{} (полиэтилен)'.format(
                    object.davlenie_name.capitalize())

            doc.tables[11].rows[1].cells[1].text = "{} МПа".format(object.prover_davl)
            if object.Data_zamera1:
                doc.tables[11].rows[1].cells[2].text = "{} - {}".format(
                    datetime.datetime.strftime(object.Data_zamera1,"%d.%m.%Y"),
                    datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
            doc.tables[11].rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.tables[11].rows[1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            remove_row(doc.tables[11],doc.tables[11].rows[1])
    else:
        remove_row(doc.tables[11],doc.tables[11].rows[1])

    start_index2 = 15
    if object.dop_dann_gazopr_v_zashit:
        for gaz_zash in object.dop_dann_gazopr_v_zashit.all():
            if gaz_zash.truba1 or gaz_zash.truba2:
                if gaz_zash.truba1.diametr or gaz_zash.truba1.x \
                    or gaz_zash.truba2.diametr or gaz_zash.truba2.x:

                    break

                else:
                    for table in doc.tables[start_index2:]:
                        for row in table.rows:
                            for cell in row.cells:
                                if 'полеЗащитныйфутляр' in cell.text:
                                    doc.tables[start_index2-2]._element.getparent().remove(
                                        doc.tables[start_index2-2]._element)
                                    doc.tables[start_index2-2]._element.getparent().remove(
                                        doc.tables[start_index2-2]._element)
                                    doc.tables[start_index2-2]._element.getparent().remove(
                                        doc.tables[start_index2-2]._element)
                                    doc.tables[start_index2-2]._element.getparent().remove(
                                        doc.tables[start_index2-2]._element)
                                    break
                        start_index2 += 1

                    for index, paragraph in enumerate(doc.paragraphs):
                        if 'прокладка газопровода в защитном футляре' in paragraph.text:
                            remove_paragraph(doc.paragraphs[index-2])
                            remove_paragraph(doc.paragraphs[index-2])
                            remove_paragraph(doc.paragraphs[index-2])

    start_index3 = 10
    if object.dop_dann_futlyar_na_vyhode:
        for fut_vyh in object.dop_dann_futlyar_na_vyhode.all():
            if fut_vyh.diametr or fut_vyh.x:
                break
            else:
                for table in doc.tables[start_index3:]:
                    for row in table.rows:
                        for cell in row.cells:
                            if 'полеФутлярнавыходе' in cell.text:

                                doc.tables[start_index3]._element.getparent().remove(
                                    doc.tables[start_index3]._element)
                                break
                    start_index3 += 1

                    for index, paragraph in enumerate(doc.paragraphs):
                        if 'установка футляров на выходе из земли' in paragraph.text:
                            remove_paragraph(doc.paragraphs[index-2])
                            remove_paragraph(doc.paragraphs[index-2])
                            remove_paragraph(doc.paragraphs[index-2])

    start_index4 = 15
    if object.dop_dann_ob_ustanovke_opor:
        for ust_opor in object.dop_dann_futlyar_na_vyhode.all():
            if ust_opor.diametr or ust_opor.x:
                break
            else:
                for table in doc.tables[start_index4:]:
                    for row in table.rows:
                        for cell in row.cells:
                            if 'полеОпораподгазопровод' in cell.text:

                                doc.tables[start_index4-2]._element.getparent().remove(
                                    doc.tables[start_index4-2]._element)
                                doc.tables[start_index4-2]._element.getparent().remove(
                                    doc.tables[start_index4-2]._element)
                                doc.tables[start_index4-2]._element.getparent().remove(
                                    doc.tables[start_index4-2]._element)
                                break
                    start_index4 += 1

                for index, paragraph in enumerate(doc.paragraphs):
                    if 'установка опор под газопровод' in paragraph.text:
                        remove_paragraph(doc.paragraphs[index-2])
                        remove_paragraph(doc.paragraphs[index-2])
                        remove_paragraph(doc.paragraphs[index-2])

    if flag_delete_svarshik2:
        remove_row(doc.tables[4], doc.tables[4].rows[5])
    if flag_delete_svarshik1:
        remove_row(doc.tables[4], doc.tables[4].rows[4])


    """
    РОСТЕХНАДЗОР
    """

    if object.rostehnadzor:

        doc_per = docx.Document('Перечень лиц, участвующих в строительстве.docx')

        style = doc_per.styles['Normal']

        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)

        style_under = doc_per.styles.add_style('Under_style', WD_STYLE_TYPE.PARAGRAPH)
        style_under.font.underline = True
        style_under.font.name = 'Times New Roman'
        style_under.font.size = Pt(10)

        for index, table in enumerate(doc_per.tables):
            for index, row in enumerate(table.rows):
                for index, cell in enumerate(row.cells):
                    if 'полеНазваниеобъекта' in cell.text:
                        if object.name_object:
                            cell.text = cell.text.replace('полеНазваниеобъекта', object.name_object)
                            cell.paragraphs[0].style = style_under
                        else:
                            cell.text = cell.text.replace('полеНазваниеобъекта', '')


        doc2 = docx.Document('Акт нивелировки.docx')

        for index, paragraph in enumerate(doc2.paragraphs):
            for index, run in enumerate(paragraph.runs):
                if 'полеДатаукладки' in run.text:
                    if object.Data_ukl:
                        run.text = run.text.replace('полеДатаукладки',
                            datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатаукладки', '')

                if 'полеНазваниеобъекта'in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта',
                            object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

                if 'полеПроектнаяорг'in run.text:
                    if object.proektnaya_org:
                        run.text = run.text.replace('полеПроектнаяорг',
                            object.proektnaya_org)
                    else:
                        run.text = run.text.replace('полеПроектнаяорг', '')

                if 'полеТехндолжность' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.post:
                                run.text = run.text.replace('полеТехндолжность',
                                    object.tehnadzor.person.post)
                            else:
                                run.text = run.text.replace('полеТехндолжность', '')
                        else:
                            run.text = run.text.replace('полеТехндолжность', '')
                    else:
                        run.text = run.text.replace('полеТехндолжность', '')


                if 'полеТехнФИО' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.fio:
                                run.text = run.text.replace('полеТехнФИО',
                                    object.tehnadzor.person.fio)
                            else:
                                run.text = run.text.replace('полеТехнФИО', '')
                        else:
                            run.text = run.text.replace('полеТехнФИО', '')
                    else:
                        run.text = run.text.replace('полеТехнФИО', '')


                if 'полеПК' in run.text:
                    if object.rostehnadzor:
                        run.text = run.text.replace('полеПК',
                            f'ПК{object.rostehnadzor.pk1}+{object.rostehnadzor.pk1_diam}-ПК{object.rostehnadzor.pk2}+{object.rostehnadzor.pk2_diam}')
                    else:
                        run.text = run.text.replace('полеПК', 'ПК_+_-ПК_+_')


        f_per = BytesIO()
        doc_per.save(f_per)

        f2 = BytesIO()
        doc2.save(f2)

        doc3 = docx.Document('Акты допом.docx')


        style = doc3.styles['Normal']

        font = style.font
        font.size = Pt(11)
        font.name = 'Times New Roman'

        style_for_name_object = doc3.styles.add_style('style_for_name_object', WD_STYLE_TYPE.PARAGRAPH)
        style_for_name_object.font.bold = True
        style_for_name_object.font.italic = True
        style_for_name_object.font.name = 'Times New Roman'
        style_for_name_object.font.size = Pt(12)

        style_for_teh_11 = doc3.styles.add_style('style_for_teh_11', WD_STYLE_TYPE.PARAGRAPH)
        style_for_teh_11.font.name = 'Times New Roman'
        style_for_teh_11.font.size = Pt(11)

        for index, table in enumerate(doc3.tables):
            for index, row in enumerate(table.rows):
                for index, cell in enumerate(row.cells):
                    if 'полеНазваниеобъекта' in cell.text:
                        if object.name_object:
                            cell.text = cell.text.replace('полеНазваниеобъекта', object.name_object)
                        else:
                            cell.text = cell.text.replace('полеНазваниеобъекта', '')

                        cell.paragraphs[0].style = style_for_name_object

                    if 'полеДатаукладки ' in cell.text:
                        if object.Data_ukl:
                            cell.text = cell.text.replace('полеДатаукладки',
                                datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                        else:
                            cell.text = cell.text.replace('полеДатаукладки', '')

                        cell.paragraphs[0].alignment = paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if 'полеТехнаддолжность' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.post:
                                    cell.text = cell.text.replace('полеТехнаддолжность',
                                        object.tehnadzor.person.post)
                                    cell.paragraphs[0].style = style_for_teh_11
                                else:
                                    cell.text = cell.text.replace('полеТехнаддолжность', '')
                            else:
                                cell.text = cell.text.replace('полеТехнаддолжность', '')
                        else:
                            cell.text = cell.text.replace('полеТехнаддолжность', '')

                    if 'полеТехнадзорФИО' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.fio:
                                    cell.text = cell.text.replace('полеТехнадзорФИО',
                                        object.tehnadzor.person.fio)
                                    cell.paragraphs[0].style = style_for_teh_11
                                else:
                                    cell.text = cell.text.replace('полеТехнадзорФИО', '')
                            else:
                                cell.text = cell.text.replace('полеТехнадзорФИО', '')
                        else:
                            cell.text = cell.text.replace('полеТехнадзорФИО', '')

                    if 'полеТехнадзор2ФИО' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.fio:
                                    cell.text = cell.text.replace('полеТехнадзор2ФИО',
                                        object.tehnadzor.person.fio)
                                    cell.paragraphs[0].style = style_for_teh_11
                                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                else:
                                    cell.text = cell.text.replace('полеТехнадзор2ФИО', '')
                            else:
                                cell.text = cell.text.replace('полеТехнадзор2ФИО', '')
                        else:
                            cell.text = cell.text.replace('полеТехнадзор2ФИО', '')


                    if 'полеНомерпроекта' in cell.text:
                        if object.Nomer_proekt:
                            cell.text = cell.text.replace('полеНомерпроекта', object.Nomer_proekt)
                        else:
                            cell.text = cell.text.replace('полеНомерпроекта', '')

                    if 'полеПроектнаяорганизация' in cell.text:
                        if object.proektnaya_org:
                            cell.text = cell.text.replace('полеПроектнаяорганизация', object.proektnaya_org)
                        else:
                            cell.text = cell.text.replace('полеПроектнаяорганизация', '')

                    if 'полеТруба' in cell.text:
                        if object.rostehnadzor:
                            if object.rostehnadzor.truba:
                                cell.text = cell.text.replace('полеТруба',
                                    f'Труба {object.rostehnadzor.truba.diametr}x{object.rostehnadzor.truba.x}')
                            else:
                                cell.text = cell.text.replace('полеТруба', '')
                        else:
                            cell.text = cell.text.replace('полеТруба', '')

                    if 'полеДатаукладки1' in cell.text:
                        if object.Data_ukl:
                            cell.text = cell.text.replace('полеДатаукладки1',
                                datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                        else:
                            cell.text = cell.text.replace('полеДатаукладки1', '')

                    if 'полеДатаукладки2' in cell.text:
                        if object.Data_ukl:
                            cell.text = cell.text.replace('полеДатаукладки2',
                                datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                        else:
                            cell.text = cell.text.replace('полеДатаукладки2','')

        f3 = BytesIO()
        doc3.save(f3)


        doc4 = docx.Document('Акт входного контроля качества.docx')

        style = doc4.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        for index, paragraph in enumerate(doc4.paragraphs):
            for index, run in enumerate(paragraph.runs):

                if 'полеДатазамера' in run.text:
                    if object.Data_zamera2:
                        run.text = run.text.replace('полеДатазамера',
                            datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатазамера', '')

                if 'полеТехнадзордолжность' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.post:
                                run.text = run.text.replace('полеТехнадзордолжность',
                                    object.tehnadzor.person.post)
                            else:
                                run.text = run.text.replace('полеТехнадзордолжность', '')
                        else:
                            run.text = run.text.replace('полеТехнадзордолжность', '')
                    else:
                        run.text = run.text.replace('полеТехнадзордолжность', '')

                if 'полеТехнадзорФИО' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.fio:
                                run.text = run.text.replace('полеТехнадзорФИО',
                                    object.tehnadzor.person.fio)
                            else:
                                run.text = run.text.replace('полеТехнадзорФИО', '')
                        else:
                            run.text = run.text.replace('полеТехнадзорФИО', '')
                    else:
                        run.text = run.text.replace('полеТехнадзорФИО', '')

                if 'полеНомерпроекта' in run.text:
                    if object.Nomer_proekt:
                        run.text = run.text.replace('полеНомерпроекта', object.Nomer_proekt)
                    else:
                        run.text = run.text.replace('полеНомерпроекта', '')

                if 'полеПроектнаяорганизация' in run.text:
                    if object.proektnaya_org:
                        run.text = run.text.replace('полеПроектнаяорганизация', object.proektnaya_org)
                    else:
                        run.text = run.text.replace('полеПроектнаяорганизация', '')

                if 'полеНазваниеобъекта' in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта', object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

        f4 = BytesIO()
        doc4.save(f4)

    if object.prodavl:

        doc5 = docx.Document('Протокол бурения.docx')

        for index, paragraph in enumerate(doc5.paragraphs):
            for index, run in enumerate(paragraph.runs):

                if 'полеНазваниеобъекта' in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта', object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

                if 'полеДиамтр' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.diametr:
                                run.text = run.text.replace('полеДиамтр', object.prodavl.truba.diametr)
                            else:
                                run.text = run.text.replace('полеДиамтр', '')
                        else:
                            run.text = run.text.replace('полеДиамтр', '')
                    else:
                        run.text = run.text.replace('полеДиамтр', '')

                if 'полеДлнтр' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.dlina:
                                run.text = run.text.replace('полеДлнтр', object.prodavl.truba.dlina)
                            else:
                                run.text = run.text.replace('полеДлнтр', '')
                        else:
                            run.text = run.text.replace('полеДлнтр', '')
                    else:
                        run.text = run.text.replace('полеДлнтр', '')

                if 'полеПК1' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_1_diam:
                            run.text = run.text.replace('полеПК1', object.prodavl.pk0_1_diam)
                        else:
                            run.text = run.text.replace('полеПК1', '')
                    else:
                        run.text = run.text.replace('полеПК1', '')

                if 'полеПК2' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_2_diam:
                            run.text = run.text.replace('полеПК2', object.prodavl.pk0_2_diam)
                        else:
                            run.text = run.text.replace('полеПК2', '')
                    else:
                        run.text = run.text.replace('полеПК2', '')

                if 'полеДатаукл' in run.text:
                    if object.Data_ukl:
                        run.text = run.text.replace('полеДатаукл',
                            datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатаукл','')

        f5 = BytesIO()
        doc5.save(f5)

        doc6 = docx.Document('Акт приёмки.docx')

        style_for_te = doc6.styles.add_style('style_for_teh_11', WD_STYLE_TYPE.PARAGRAPH)
        style_for_te.font.name = 'Times New Roman'
        style_for_te.font.size = Pt(11)


        for index, paragraph in enumerate(doc6.paragraphs):
            for index, run in enumerate(paragraph.runs):

                if 'полеДатаукладки' in run.text:
                    if object.Data_ukl:
                        run.text = run.text.replace('полеДатаукладки',
                            datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатаукладки','')

                if 'полеНазваниеобъекта' in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта', object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

                if 'полеТехнФИО' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.fio:
                                run.text = run.text.replace('полеТехнФИО',
                                    object.tehnadzor.person.fio)
                            else:
                                run.text = run.text.replace('полеТехнФИО', '')
                        else:
                            run.text = run.text.replace('полеТехнФИО', '')
                    else:
                        run.text = run.text.replace('полеТехнФИО', '')

                if 'полеТрх' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.x:
                                run.text = run.text.replace('полеТрх', object.prodavl.truba.x)
                            else:
                                run.text = run.text.replace('полеТрх', '')
                        else:
                            run.text = run.text.replace('полеТрх', '')
                    else:
                        run.text = run.text.replace('полеТрх', '')


                if 'полеТрдиам' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.diametr:
                                run.text = run.text.replace('полеТрдиам', object.prodavl.truba.diametr)
                            else:
                                run.text = run.text.replace('полеТрдиам', '')
                        else:
                            run.text = run.text.replace('полеТрдиам', '')
                    else:
                        run.text = run.text.replace('полеТрдиам', '')

                if 'полеТрдл' in run.text:
                    if object.prodavl:
                        if object.prodavl.truba:
                            if object.prodavl.truba.dlina:
                                run.text = run.text.replace('полеТрдл', object.prodavl.truba.dlina)
                            else:
                                run.text = run.text.replace('полеТрдл', '')
                        else:
                            run.text = run.text.replace('полеТрдл', '')
                    else:
                        run.text = run.text.replace('полеТрдл', '')

                if 'полеПК1' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_1_diam:
                            run.text = run.text.replace('полеПК1', object.prodavl.pk0_1_diam)
                        else:
                            run.text = run.text.replace('полеПК1', '')
                    else:
                        run.text = run.text.replace('полеПК1', '')

                if 'полеПК2' in run.text:
                    if object.prodavl:
                        if object.prodavl.pk0_2_diam:
                            run.text = run.text.replace('полеПК2', object.prodavl.pk0_2_diam)
                        else:
                            run.text = run.text.replace('полеПК2', '')
                    else:
                        run.text = run.text.replace('полеПК2', '')

                if 'полеПроектнаяорг' in run.text:
                    if object.proektnaya_org:
                        run.text = run.text.replace('полеПроектнаяорг', object.proektnaya_org)
                    else:
                        run.text = run.text.replace('полеПроектнаяорг', '')

                if 'полеНомерпроекта' in run.text:
                    if object.Nomer_proekt:
                        run.text = run.text.replace('полеНомерпроекта', object.Nomer_proekt)
                    else:
                        run.text = run.text.replace('полеНомерпроекта', '')

        for index, table in enumerate(doc6.tables):
            for index, row in enumerate(table.rows):
                for index, cell in enumerate(row.cells):
                    if 'полеТехнФИО' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.fio:
                                    cell.text = cell.text.replace('полеТехнФИО',
                                        object.tehnadzor.person.fio)
                                    cell.paragraphs[0].style = style_for_te
                                else:
                                    cell.text = cell.text.replace('полеТехнФИО', '')
                            else:
                                cell.text = cell.text.replace('полеТехнФИО', '')
                        else:
                            cell.text = cell.text.replace('полеТехнФИО', '')

                    if 'полеТехнаддолж' in cell.text:
                        if object.tehnadzor:
                            if object.tehnadzor.person:
                                if object.tehnadzor.person.post:
                                    cell.text = cell.text.replace('полеТехнаддолж',
                                        object.tehnadzor.person.post)
                                    cell.paragraphs[0].style = style_for_te
                                else:
                                    cell.text = cell.text.replace('полеТехнаддолж', '')
                            else:
                                cell.text = cell.text.replace('полеТехнаддолж', '')
                        else:
                            cell.text = cell.text.replace('полеТехнаддолж', '')

        f6 = BytesIO()
        doc6.save(f6)

    if object.rostehnadzor and object.prodavl:

        response = HttpResponse(content_type='application/zip')
        zf = ZipFile(response, 'w')

        f = BytesIO()
        doc.save(f)

        zf.writestr('ИТД.docx', f.getvalue())
        zf.writestr('Перечень лиц, участвующих в строительстве.docx', f_per.getvalue())
        zf.writestr('Акт нивелировки.docx', f2.getvalue())
        zf.writestr('Акты допом.docx', f3.getvalue())
        zf.writestr('Акт входного контроля качества.docx', f4.getvalue())
        zf.writestr('Протокол бурения.docx', f5.getvalue())
        zf.writestr('Акт приёмки.docx', f6.getvalue())

        response['Content-Disposition'] = 'attachment; filename=ITD.zip'
        return response

    elif object.rostehnadzor:

        response = HttpResponse(content_type='application/zip')
        zf = ZipFile(response, 'w')

        f = BytesIO()
        doc.save(f)

        zf.writestr('ИТД.docx', f.getvalue())
        zf.writestr('Перечень лиц, участвующих в строительстве.docx', f_per.getvalue())
        zf.writestr('Акт нивелировки.docx', f2.getvalue())
        zf.writestr('Акты допом.docx', f3.getvalue())
        zf.writestr('Акт входного контроля качества.docx', f4.getvalue())

        response['Content-Disposition'] = 'attachment; filename=ITD.zip'
        return response

    elif object.prodavl:

        response = HttpResponse(content_type='application/zip')
        zf = ZipFile(response, 'w')

        f = BytesIO()
        doc.save(f)

        zf.writestr('ИТД.docx', f.getvalue())
        zf.writestr('Протокол бурения.docx', f5.getvalue())
        zf.writestr('Акт приёмки.docx', f6.getvalue())

        response['Content-Disposition'] = 'attachment; filename=ITD.zip'
        return response

    else:

        f = BytesIO()
        doc.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = "attachment; filename='ITD.docx'"
        response['Content-Length'] = length

        return response

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
