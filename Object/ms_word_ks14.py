import datetime
from io import BytesIO

from django.http import HttpResponse
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from Object.models import *


object_fields_str = {'полеМестоположение':'place','полеПриказ':'order',\
'полеНазваниеобъекта':'name_object', 'полеКодобъекта':'kod_object', \
'полеНомерпроекта':'Nomer_proekt', 'полеНомердоговоранаразм':'Nomer_razm',\
'полеПроектнаяорганизация':'proektnaya_org','полеКонтрагент':'full_kontragent',
'полеНомерзаданиянапроект':'Nomer_zadaniya'}

object_fields_date = {'полеДатадоговоранаразм':'Data_razm',\
 'полеДатазаданиянапроект':'Date_proekt', 'полеДатаначаларабот':'date_nachala_rabot',
 'полеДатасосткс2':''}

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def form_ks14(name):
    object = Object.objects.get(name_object=name)

    object_fields = {}

    if object.kontragent:
        if object.kontragent.name_kontragent:
            object_fields['полеНазваниеконтрагента'] = object.kontragent.name_kontragent
        else:
            object_fields['полеНазваниеконтрагента'] = ''

        if object.kontragent.podpisant:
            object_fields['полеПредставительгенподрядчика'] = object.kontragent.podpisant.post
            object_fields['полеПредставительгенподрядчФИО'] = object.kontragent.podpisant.fio
        else:
            object_fields['полеПредставительгенподрядчика'] = ''
            object_fields['полеПредставительгенподрядчФИО'] = ''
    else:
        object_fields['полеПредставительгенподрядчика'] = ''
        object_fields['полеПредставительгенподрядчФИО'] = ''
        object_fields['полеНазваниеконтрагента'] = ''

    if object.ks11_predsedatel:

        object_fields['полеПредседателькомиссии'] = object.ks11_predsedatel.post
        object_fields['полеПредседателькомиссФИО'] = object.ks11_predsedatel.fio

    else:

        object_fields['полеПредседателькомиссии'] = ''
        object_fields['полеПредседателькомиссФИО'] = ''


    if object.ks11_predstav_proekt:

        object_fields['полеПредставительпроектир'] = object.ks11_predstav_proekt.post
        object_fields['полеПредставительпроектФИО'] = object.ks11_predstav_proekt.fio

    else:

        object_fields['полеПредставительпроектир'] = ''
        object_fields['полеПредставительпроектФИО'] = ''

    if object.ks11_predstav_ekspl:

        object_fields['полеПредставительэкспл'] = object.ks11_predstav_ekspl.post
        object_fields['полеПредставительэксФИО'] = object.ks11_predstav_ekspl.fio

    else:

        object_fields['полеПредставительэкспл'] = ''
        object_fields['полеПредставительэксФИО'] = ''

    if object.smeta:
        if object.smeta.date_nach_zakr:
            object_fields['полеДатаначалазакрытия'] = datetime.date.strftime(object.smeta.date_nach_zakr,"%d.%m.%Y")
        else:
            object_fields['полеДатаначалазакрытия'] = ''

        if object.smeta.date_kon_zakr:
            object_fields['полеДатаконцазакрытия'] = datetime.date.strftime(object.smeta.date_kon_zakr,"%d.%m.%Y")
        else:
            object_fields['полеДатаконцазакрытия'] = ''
    else:
        object_fields['полеДатаконцазакрытия'] = ''
        object_fields['полеДатаначалазакрытия'] = ''

    doc = docx.Document('KS11_KS14_template.docx')

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    style_bold = doc.styles.add_style('Bold_style', WD_STYLE_TYPE.PARAGRAPH)
    style_bold.font.bold = True
    style_bold.font.size = Pt(10)

    style_11 = doc.styles.add_style('Normal_11', WD_STYLE_TYPE.PARAGRAPH)
    style_11.font.italic = True
    style_11.font.bold = True
    style_11.font.size = Pt(10)

    style_10 = doc.styles.add_style('Normal_10', WD_STYLE_TYPE.PARAGRAPH)
    style_10.font.size = Pt(10)

    new_paragraph = doc.add_paragraph()
    new_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:

            """
            ГАЗОПРОВОДЫ
            """

            if 'Списокгазопроводов' in run.text:

                paragraph.text = ''

                paragraph.add_run("{}\n".format(object.name_gazoprovod)).bold = True

                if object.gazoprovod_podzem_stal:
                    if object.gazoprovod_podzem_stal.truba:
                        paragraph.add_run("{}\n".format(object.name_podzem_stal_gazoprovod)).bold = True
                        for truba in object.gazoprovod_podzem_stal.truba.all():
                            if truba.diametr:
                                paragraph.add_run("Труба Ø{}x{} - {} ({})\n".format(truba.diametr,
                                   truba.x,
                                   truba.dlina,
                                   truba.prim))

                    paragraph.add_run("Установлено:\n").bold = True

                    if object.gazoprovod_podzem_stal.neraz_soed:
                        if object.gazoprovod_podzem_stal.neraz_soed.PE:
                            paragraph.add_run("Неразъемное соединение ПЭØ{}/СТØ{} - {} шт.\n".format(
                                object.gazoprovod_podzem_stal.neraz_soed.PE,
                                object.gazoprovod_podzem_stal.neraz_soed.ST,
                                object.gazoprovod_podzem_stal.neraz_soed.kolvo))

                    if object.gazoprovod_podzem_stal.kontrolnaya_trubka:
                        paragraph.add_run("Контрольная трубка - {} шт.\n".format(
                            object.gazoprovod_podzem_stal.kontrolnaya_trubka))

                    if object.gazoprovod_podzem_stal.otvod_90:
                        paragraph.add_run("Отвод ст. 90 гр. - {} шт.\n".format(
                            object.gazoprovod_podzem_stal.otvod_90))

                    if object.gazoprovod_podzem_stal.opoznavat_znak:
                        paragraph.add_run("Опознавательный знак - {} шт.".format(
                            object.gazoprovod_podzem_stal.opoznavat_znak))

                if object.gazoprovod_podzem_poliet:
                    if object.gazoprovod_podzem_poliet.truba:
                        paragraph.add_run("\n{}\n".format(object.name_podzem_poliet_gazoprovod)).bold = True
                        try:
                            for truba in object.gazoprovod_podzem_poliet.truba.all():
                                if truba.diametr:
                                    paragraph.add_run("Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                                        truba.diametr, truba.x, truba.dlina))
                        except:
                            for truba in object.gazoprovod_podzem_poliet.truba.all():
                                if truba.diametr:
                                    paragraph.add_run("Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                                        truba.diametr, truba.x, truba.dlina))

                    paragraph.add_run("Установлено:\n").bold = True

                    if object.gazoprovod_podzem_poliet.mufta:
                        for mufta in object.gazoprovod_podzem_poliet.mufta.all():
                            if mufta.diametr:
                                paragraph.add_run("Муфта электросварная ПЭ100 SDR11 Ø{}х{} – {} шт.\n".format(
                                    mufta.diametr, mufta.x, mufta.dlina))

                    if object.gazoprovod_podzem_poliet.otvod:
                        if object.gazoprovod_podzem_poliet.otvod.diametr:
                            paragraph.add_run("Отвод ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.otvod.diametr,
                                object.gazoprovod_podzem_poliet.otvod.kolvo))

                    if object.gazoprovod_podzem_poliet.troinik:
                        if object.gazoprovod_podzem_poliet.troinik.diametr1:
                            paragraph.add_run("Тройник электросварной ПЭ100 SDR11 Ø{}xØ{}xØ{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.troinik.diametr1,
                                object.gazoprovod_podzem_poliet.troinik.diametr2,
                                object.gazoprovod_podzem_poliet.troinik.diametr3,
                                object.gazoprovod_podzem_poliet.troinik.kolvo))

                    if object.gazoprovod_podzem_poliet.zaglushka:
                        if object.gazoprovod_podzem_poliet.zaglushka.diametr:
                            paragraph.add_run("Заглушка электросварная ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.zaglushka.diametr,
                                object.gazoprovod_podzem_poliet.zaglushka.kolvo))

                    if object.gazoprovod_podzem_poliet.lenta:
                        paragraph.add_run("Лента сигнальная “Осторожно! ГАЗ!”  – {} м\n".format(
                            object.gazoprovod_podzem_poliet.lenta))

                    if object.gazoprovod_podzem_poliet.kran:
                        if object.gazoprovod_podzem_poliet.kran.diametr:
                            paragraph.add_run("Кран шаровой для подземной установки ПЭ 100 ГАЗ SDR11 Ø{} – {} шт.\n".format(
                                object.gazoprovod_podzem_poliet.kran.diametr,
                                object.gazoprovod_podzem_poliet.kran.kolvo))

                    if object.gazoprovod_podzem_poliet.znak:
                        paragraph.add_run("Опознавательный знак – {} шт.\n".format(
                            object.gazoprovod_podzem_poliet.znak))

                    if object.gazoprovod_podzem_poliet.sedelka:
                        if object.gazoprovod_podzem_poliet.sedelka.diametr1:
                            paragraph.add_run("Седелка электросварная Ø{}xØ{}xØ{} – {} шт.\n".format(
                            object.gazoprovod_podzem_poliet.sedelka.diametr1,
                            object.gazoprovod_podzem_poliet.sedelka.diametr2,
                            object.gazoprovod_podzem_poliet.sedelka.diametr3,
                            object.gazoprovod_podzem_poliet.sedelka.kolvo))

                if object.gazoprovod_nadzem_stal:
                    paragraph.add_run("{}\n".format(object.name_nadzem_stal_gazoprovod)).bold = True

                    if object.gazoprovod_nadzem_stal.truba:
                        for truba in object.gazoprovod_nadzem_stal.truba.all():
                            if truba.diametr:
                                paragraph.add_run("Труба стальная Ø{}х{} – {} м\n".format(
                                    truba.diametr, truba.x, truba.dlina))

                    paragraph.add_run("Установлено:\n").bold = True

                    if object.gazoprovod_nadzem_stal.izolir_soed:
                        if object.gazoprovod_nadzem_stal.izolir_soed.diametr:
                            paragraph.add_run("Изолирующие соединения Ø{} – {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.izolir_soed.diametr,
                                object.gazoprovod_nadzem_stal.izolir_soed.kolvo))

                    if object.gazoprovod_nadzem_stal.kran_stal:
                        if object.gazoprovod_nadzem_stal.kran_stal.diametr:
                            paragraph.add_run("Кран стальной Ø{} – {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.kran_stal.diametr,
                                object.gazoprovod_nadzem_stal.kran_stal.kolvo))

                    if object.gazoprovod_nadzem_stal.otvod:
                        if object.gazoprovod_nadzem_stal.otvod.diametr:
                            paragraph.add_run("Отвод Ø{} – {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.otvod.diametr,
                                object.gazoprovod_nadzem_stal.otvod.kolvo))

                    if object.gazoprovod_nadzem_stal.cokolnyi_vvod:
                        if object.gazoprovod_nadzem_stal.cokolnyi_vvod.PE:
                            paragraph.add_run("Цокольный ввод ПЭØ{}/СТØ{} - {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.cokolnyi_vvod.PE,
                                object.gazoprovod_nadzem_stal.cokolnyi_vvod.ST,
                                object.gazoprovod_nadzem_stal.cokolnyi_vvod.kolvo))

                    if object.gazoprovod_nadzem_stal.kreplenie:
                        paragraph.add_run("Крепление газопровода к кирпичной стене – {} шт.\n".format(
                            object.gazoprovod_nadzem_stal.kreplenie))

                    if object.gazoprovod_nadzem_stal.stoika:
                        if object.gazoprovod_nadzem_stal.stoika.diametr:
                            paragraph.add_run("Стойка под газопровод Ø{} – {} м - {} шт.\n".format(
                                object.gazoprovod_nadzem_stal.stoika.diametr,
                                object.gazoprovod_nadzem_stal.stoika.dlina,
                                object.gazoprovod_nadzem_stal.stoika.kolvo))

                if object.itogo != 0.0:
                    paragraph.add_run("Итого протяженность газопровода – {} м".format(
                        object.itogo)).bold = True

                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.style = style_10

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field in object_fields:
                    if field in cell.text:
                        if object_fields[field]:
                            cell.text = cell.text.replace(field,object_fields[field])
                            cell.paragraphs[0].style = style_11
                        else:
                            cell.text = cell.text.replace(field, '')


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field in object_fields_str:
                    if field in cell.text:
                        if getattr(object,object_fields_str[field]):
                            cell.text = cell.text.replace(field,getattr(object,object_fields_str[field]))
                            cell.paragraphs[0].style = style_11
                        else:
                            cell.text = cell.text.replace(field, '')


    for paragraph in doc.paragraphs:
        for field in object_fields_str.keys():
            if field in paragraph.text:
                if getattr(object,object_fields_str[field]):
                    paragraph.text = paragraph.text.replace(field,getattr(object,object_fields_str[field]))
                    paragraph.style = style_11
                else:
                    paragraph.text = paragraph.text.replace(field, '')



        for field in object_fields_date.keys():
            if field in paragraph.text:
                if field == "полеДатаначаларабот":
                    if object.smeta:
                        if object.smeta.date_nach_rabot:
                            paragraph.text = paragraph.text.replace(field, object.date_nachala_rabot)
                            paragraph.style = style_11
                        else:
                            paragraph.text = paragraph.text.replace(field, '')
                    else:
                        paragraph.text = paragraph.text.replace(field, '')

                elif field == "полеДатасосткс2":
                    if object.smeta:
                        if object.smeta.date_ks2:
                            paragraph.text = paragraph.text.replace(field, object.date_ks2_2)
                            paragraph.style = style_11
                        else:
                            paragraph.text = paragraph.text.replace(field, '')
                    else:
                        paragraph.text = paragraph.text.replace(field, '')

                else:
                    try:
                        paragraph.text = paragraph.text.replace(field,datetime.date.strftime(getattr(object,object_fields_date[field]),"%d.%m.%Y"))
                        paragraph.style = style_11
                    except:
                        paragraph.text = paragraph.text.replace(field, '')


    flag_gazoprovod_podzem_stal = False
    flag_gazoprovod_podzem_poliet = False
    flag_gazoprovod_nadzem_stal = False

    sum1 = 0.0

    if object.gazoprovod_nadzem_stal:
        if object.gazoprovod_nadzem_stal.truba:
            for truba in object.gazoprovod_nadzem_stal.truba.all():
                if truba.diametr:
                    doc.tables[2].rows[5].cells[0].text = doc.tables[2].rows[5].cells[0].text.replace('полеНадземсталгазопр', object.full_name_nadzem_stal_gazoprovod)
                    sum1 += float(object.dlina_nadzem_stal)
                    doc.tables[2].rows[5].cells[2].text = str(object.dlina_nadzem_stal)
                    doc.tables[2].rows[5].cells[4].text = str(object.dlina_nadzem_stal)
                    doc.tables[2].rows[5].cells[0].paragraphs[0].style = style_11
                    doc.tables[2].rows[5].cells[2].paragraphs[0].style = style_11
                    doc.tables[2].rows[5].cells[4].paragraphs[0].style = style_11
                    doc.tables[2].rows[5].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.tables[2].rows[5].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
                else:
                    remove_row(doc.tables[2], doc.tables[2].rows[5])
        else:
            remove_row(doc.tables[2], doc.tables[2].rows[5])
    else:
        remove_row(doc.tables[2], doc.tables[2].rows[5])

    if object.gazoprovod_podzem_poliet:
        if object.gazoprovod_podzem_poliet.truba:
            for truba in object.gazoprovod_podzem_poliet.truba.all():
                if truba.diametr:
                    doc.tables[2].rows[4].cells[0].text = doc.tables[2].rows[4].cells[0].text.replace('полеПодземполиетгазопр', object.full_name_podzem_poliet_gazoprovod)
                    sum1 += float(object.dlina_podzem_poliet)
                    doc.tables[2].rows[4].cells[2].text = str(object.dlina_podzem_poliet)
                    doc.tables[2].rows[4].cells[4].text = str(object.dlina_podzem_poliet)
                    doc.tables[2].rows[4].cells[0].paragraphs[0].style = style_11
                    doc.tables[2].rows[4].cells[2].paragraphs[0].style = style_11
                    doc.tables[2].rows[4].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.tables[2].rows[4].cells[4].paragraphs[0].style = style_11
                    doc.tables[2].rows[4].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
                else:
                    remove_row(doc.tables[2], doc.tables[2].rows[4])
        else:
            remove_row(doc.tables[2], doc.tables[2].rows[4])
    else:
        remove_row(doc.tables[2], doc.tables[2].rows[4])


    if object.gazoprovod_podzem_stal:
        if object.gazoprovod_podzem_stal.truba:
            for truba in object.gazoprovod_podzem_stal.truba.all():
                if truba.diametr:
                    doc.tables[2].rows[3].cells[0].text = doc.tables[2].rows[3].cells[0].text.replace('полеПодземсталгазопр', object.full_name_podzem_stal_gazoprovod)
                    sum1 += float(object.dlina_podzem_stal)
                    doc.tables[2].rows[3].cells[2].text = str(object.dlina_podzem_stal)
                    doc.tables[2].rows[3].cells[4].text = str(object.dlina_podzem_stal)
                    doc.tables[2].rows[3].cells[0].paragraphs[0].style = style_11
                    doc.tables[2].rows[3].cells[2].paragraphs[0].style = style_11
                    doc.tables[2].rows[3].cells[4].paragraphs[0].style = style_11
                    doc.tables[2].rows[3].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.tables[2].rows[3].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
                else:
                    remove_row(doc.tables[2], doc.tables[2].rows[3])
        else:
            remove_row(doc.tables[2], doc.tables[2].rows[3])
    else:
        remove_row(doc.tables[2], doc.tables[2].rows[3])

    for row in doc.tables[2].rows:
        for cell in row.cells:
            if 'Мощность' in cell.text:
                row.cells[2].text = str(sum1)
                row.cells[4].text = str(sum1)
                row.cells[0].paragraphs[0].style = style_11
                row.cells[2].paragraphs[0].style = style_11
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[4].paragraphs[0].style = style_11
                row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    summa_ks2_bez_nds_rub = 0
    summa_ks2_bez_nds_kop = 0

    if object.smeta:
        if object.smeta.summa_ks2_bez_nds:
            if '.' in str(object.smeta.summa_ks2_bez_nds):
                try:
                    summa_ks2_bez_nds_rub = object.smeta.summa_ks2_bez_nds.split('.')[0]
                    summa_ks2_bez_nds_kop = object.smeta.summa_ks2_bez_nds.split('.')[1]
                except:
                    summa_ks2_bez_nds_rub = ''
                    summa_ks2_bez_nds_kop = ''

            if ',' in str(object.smeta.summa_ks2_bez_nds):
                try:
                    summa_ks2_bez_nds_rub = object.smeta.summa_ks2_bez_nds.split(',')[0]
                    summa_ks2_bez_nds_kop = object.smeta.summa_ks2_bez_nds.split(',')[1]
                except:
                    summa_ks2_bez_nds_rub = ''
                    summa_ks2_bez_nds_kop = ''

            else:
                summa_ks2_bez_nds_rub = object.smeta.summa_ks2_bez_nds
                summa_ks2_bez_nds_kop = 0

        else:
            summa_ks2_bez_nds_rub = 0
            summa_ks2_bez_nds_kop = 0
    else:
        summa_ks2_bez_nds_rub = 0
        summa_ks2_bez_nds_kop = 0

    if summa_ks2_bez_nds_rub:
        doc.tables[4].rows[0].cells[1].text = doc.tables[4].rows[0].cells[1].text.replace('полеСуммабезндс', str(summa_ks2_bez_nds_rub))
        doc.tables[4].rows[2].cells[2].text = doc.tables[4].rows[2].cells[2].text.replace('полеСуммабезндс', str(summa_ks2_bez_nds_rub))
        doc.tables[4].rows[0].cells[1].paragraphs[0].style = style_11
        doc.tables[4].rows[2].cells[2].paragraphs[0].style = style_11
    else:
        doc.tables[4].rows[0].cells[1].text = doc.tables[4].rows[0].cells[1].text.replace('полеСуммабезндс', '0')
        doc.tables[4].rows[2].cells[2].text = doc.tables[4].rows[2].cells[2].text.replace('полеСуммабезндс', '0')
        doc.tables[4].rows[0].cells[1].paragraphs[0].style = style_11
        doc.tables[4].rows[2].cells[2].paragraphs[0].style = style_11

    if summa_ks2_bez_nds_kop:
        doc.tables[4].rows[0].cells[5].text = doc.tables[4].rows[0].cells[5].text.replace('кк', str(summa_ks2_bez_nds_kop))
        doc.tables[4].rows[2].cells[5].text = doc.tables[4].rows[2].cells[5].text.replace('кк', str(summa_ks2_bez_nds_kop))
        doc.tables[4].rows[0].cells[5].paragraphs[0].style = style_11
        doc.tables[4].rows[2].cells[5].paragraphs[0].style = style_11
        doc.tables[4].rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[4].rows[2].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.tables[4].rows[0].cells[5].text = doc.tables[4].rows[0].cells[5].text.replace('кк', '0')
        doc.tables[4].rows[2].cells[5].text = doc.tables[4].rows[2].cells[5].text.replace('кк', '0')
        doc.tables[4].rows[0].cells[5].paragraphs[0].style = style_11
        doc.tables[4].rows[2].cells[5].paragraphs[0].style = style_11
        doc.tables[4].rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[4].rows[2].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    f = BytesIO()
    doc.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = "attachment; filename='KS11_KS14.docx'"
    response['Content-Length'] = length
    
    return response
