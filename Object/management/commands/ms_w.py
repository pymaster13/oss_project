import datetime

from django.core.management.base import BaseCommand
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from Object.models import *

object_fields_str = {'полеМестоположение':'place','полеПриказ':'order',\
'полеНазваниеобъекта':'name_object', 'полеКодобъекта':'kod_object', \
'полеНомерпроекта':'Nomer_proekt', 'полеНомердоговоранаразм':'Nomer_razm',\
'полеПроектнаяорганизация':'proektnaya_org','полеКонтрагент':'full_kontragent',
'полеНомерзаданиянапроект':'Nomer_zadaniya'}

object_fields_date = {'полеДатадоговоранаразм':'Data_razm',\
 'полеДатазаданиянапроект':'Date_proekt', 'полеДатаначаларабот':'date_nachala_rabot'}

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

class Command(BaseCommand):
    help = 'test MS_WORD'

    def handle(self, *args, **options):

        object = Object.objects.first()

        object_fields = {'полеПредседателькомиссии':object.ks11_predsedatel.post,\
        'полеПредседателькомиссФИО':object.ks11_predsedatel.fio, \
        'полеПредставительпроектир':object.ks11_predstav_proekt.post,\
        'полеПредставительпроектФИО':object.ks11_predstav_proekt.fio,\
        'полеПредставительгенподрядчика':object.kontragent.podpisant.post,\
        'полеПредставительгенподрядчФИО':object.kontragent.podpisant.fio,\
        'полеПредставительэкспл':object.ks11_predstav_ekspl.post,\
        'полеПредставительэксФИО':object.ks11_predstav_ekspl.fio,\
        'полеДатаначалазакрытия':datetime.date.strftime(object.smeta.date_nach_zakr,"%d.%m.%Y"),\
        'полеДатаконцазакрытия':datetime.date.strftime(object.smeta.date_kon_zakr,"%d.%m.%Y")}

        doc = docx.Document('KS11_KS14_template.docx')

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        style_bold = doc.styles.add_style('Bold_style', WD_STYLE_TYPE.PARAGRAPH)
        style_bold.font.bold = True
        style_bold.font.size = Pt(11)

        style_11 = doc.styles.add_style('Normal_11', WD_STYLE_TYPE.PARAGRAPH)
        style_11.font.size = Pt(11)

        new_paragraph = doc.add_paragraph()
        new_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:

                """
                ГАЗОПРОВОДЫ
                """

                if 'Списокгазопроводов' in run.text:

                    paragraph.text = ''

                    paragraph.add_run("{}\n\n".format(object.name_gazoprovod)).bold = True

                    paragraph.add_run("{}\n".format(object.name_podzem_stal_gazoprovod)).bold = True
                    for truba in object.gazoprovod_podzem_stal.truba.all():
                        paragraph.add_run("- Труба Ø{}x{} - {} ({})\n".format(truba.diametr,
                                                                       truba.x,
                                                                       truba.dlina,
                                                                        truba.prim))
                    paragraph.add_run("  Установлено:\n")

                    paragraph.add_run("    - Неразъемное соединение ПЭØ{}/СТØ{} - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.neraz_soed.PE,
                        object.gazoprovod_podzem_stal.neraz_soed.ST,
                        object.gazoprovod_podzem_stal.neraz_soed.kolvo))

                    paragraph.add_run("    - Контрольная трубка - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.kontrolnaya_trubka))

                    paragraph.add_run("    - Отвод ст. 90 гр. - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.otvod_90))

                    paragraph.add_run("    - Опознавательный знак - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.opoznavat_znak))

                    paragraph.add_run("\n{}\n".format(object.name_podzem_poliet_gazoprovod)).bold = True

                    try:
                        for truba in object.gazoprovod_podzem_poliet.truba.all():
                            paragraph.add_run("- Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                                truba.diametr, truba.x, truba.dlina))
                    except:
                        for truba in object.gazoprovod_podzem_poliet.truba.all():
                            paragraph.add_run("- Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                            truba.diametr, truba.x, truba.dlina))


                    paragraph.add_run("  Установлено:\n")

                    for mufta in object.gazoprovod_podzem_poliet.mufta.all():
                        paragraph.add_run("    - Муфта электросварная ПЭ100 SDR11 Ø{}х{} – {} шт.\n".format(
                            mufta.diametr, mufta.x, mufta.dlina))

                    paragraph.add_run("    - Отвод ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.otvod.diametr,
                        object.gazoprovod_podzem_poliet.otvod.kolvo))

                    paragraph.add_run("    - Тройник электросварной ПЭ100 SDR11 Ø{}xØ{}xØ{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.troinik.diametr1,
                        object.gazoprovod_podzem_poliet.troinik.diametr2,
                        object.gazoprovod_podzem_poliet.troinik.diametr3,
                        object.gazoprovod_podzem_poliet.troinik.kolvo))

                    paragraph.add_run("    - Заглушка электросварная ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.zaglushka.diametr,
                        object.gazoprovod_podzem_poliet.zaglushka.kolvo))

                    paragraph.add_run("    - Лента сигнальная “Осторожно! ГАЗ!”  – {} м\n".format(
                        object.gazoprovod_podzem_poliet.lenta))

                    paragraph.add_run("    - Кран шаровой для подземной установки ПЭ 100 ГАЗ SDR11 Ø{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.kran.diametr,
                        object.gazoprovod_podzem_poliet.kran.kolvo))

                    paragraph.add_run("    - Опознавательный знак – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.znak))

                    paragraph.add_run("    - Седелка электросварная Ø{}xØ{}xØ{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.sedelka.diametr1,
                        object.gazoprovod_podzem_poliet.sedelka.diametr2,
                        object.gazoprovod_podzem_poliet.sedelka.diametr3,
                        object.gazoprovod_podzem_poliet.sedelka.kolvo))

                    paragraph.add_run("\n{}\n".format(object.name_nadzem_stal_gazoprovod)).bold = True

                    for truba in object.gazoprovod_nadzem_stal.truba.all():
                        paragraph.add_run("- Труба стальная Ø{}х{} – {} м\n".format(
                            truba.diametr, truba.x, truba.dlina))

                    paragraph.add_run("  Установлено:\n")

                    paragraph.add_run("    - Изолирующие соединения Ø{} – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.izolir_soed.diametr,
                        object.gazoprovod_nadzem_stal.izolir_soed.kolvo))

                    paragraph.add_run("    - Кран стальной Ø{} – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.kran_stal.diametr,
                        object.gazoprovod_nadzem_stal.kran_stal.kolvo))

                    paragraph.add_run("    - Отвод Ø{} – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.otvod.diametr,
                        object.gazoprovod_nadzem_stal.otvod.kolvo))

                    paragraph.add_run("    - Цокольный ввод ПЭØ{}/СТØ{} - {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.cokolnyi_vvod.PE,
                        object.gazoprovod_nadzem_stal.cokolnyi_vvod.ST,
                        object.gazoprovod_nadzem_stal.cokolnyi_vvod.kolvo))

                    paragraph.add_run("    - Крепление газопровода к кирпичной стене – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.kreplenie))

                    paragraph.add_run("    - Стойка под газопровод Ø{} – {} м - {} шт.\n\n".format(
                        object.gazoprovod_nadzem_stal.stoika.diametr,
                        object.gazoprovod_nadzem_stal.stoika.dlina,
                        object.gazoprovod_nadzem_stal.stoika.kolvo))

                    paragraph.add_run("Итого протяженность газопровода – {} м\n".format(
                        object.itogo)).bold = True

                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.style = style_11

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field in object_fields:
                        if field in cell.text:
                            cell.text = cell.text.replace(field,object_fields[field])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field in object_fields_str:
                        if field in cell.text:
                            cell.text = cell.text.replace(field,getattr(object,object_fields_str[field]))

        for paragraph in doc.paragraphs:
            for field in object_fields_str.keys():
                if field in paragraph.text:
                    paragraph.text = paragraph.text.replace(field,getattr(object,object_fields_str[field]))


            for field in object_fields_date.keys():
                if field in paragraph.text:
                    if field != "полеДатаначаларабот":
                        try:
                            paragraph.text = paragraph.text.replace(field,datetime.date.strftime(getattr(object,object_fields_date[field]),"%d.%m.%Y"))
                        except:
                            pass
                    else:
                        paragraph.text = paragraph.text.replace(field,object.date_nachala_rabot)

        flag_gazoprovod_podzem_stal = False
        flag_gazoprovod_podzem_poliet = False
        flag_gazoprovod_nadzem_stal = False
        if object.gazoprovod_podzem_stal:
            doc.tables[2].rows[3].cells[0].text = doc.tables[2].rows[3].cells[0].text.replace('полеПодземсталгазопр', object.full_name_podzem_stal_gazoprovod)
            doc.tables[2].rows[3].cells[2].text = str(object.dlina_podzem_stal)
            doc.tables[2].rows[3].cells[4].text = str(object.dlina_podzem_stal)
        else:
            flag_gazoprovod_podzem_stal = True
        if object.gazoprovod_podzem_poliet:
            doc.tables[2].rows[4].cells[0].text = doc.tables[2].rows[4].cells[0].text.replace('полеПодземполиетгазопр', object.full_name_podzem_poliet_gazoprovod)
            doc.tables[2].rows[4].cells[2].text = str(object.dlina_podzem_poliet)
            doc.tables[2].rows[4].cells[4].text = str(object.dlina_podzem_poliet)
        else:
            flag_gazoprovod_podzem_poliet = True
        if object.gazoprovod_nadzem_stal:
            doc.tables[2].rows[5].cells[0].text = doc.tables[2].rows[5].cells[0].text.replace('полеНадземсталгазопр', object.full_name_nadzem_stal_gazoprovod)
            doc.tables[2].rows[5].cells[2].text = str(object.dlina_nadzem_stal)
            doc.tables[2].rows[5].cells[4].text = str(object.dlina_nadzem_stal)
        else:
            flag_gazoprovod_nadzem_stal = True
        if flag_gazoprovod_podzem_stal:
            remove_row(doc.tables[2], doc.tables[2].rows[4])
        if flag_gazoprovod_podzem_poliet:
            remove_row(doc.tables[2], doc.tables[2].rows[5])
        if flag_gazoprovod_nadzem_stal:
            remove_row(doc.tables[2], doc.tables[2].rows[6])

        summa_ks2_bez_nds_rub = 0
        summa_ks2_bez_nds_kop = 0

        if object.smeta.summa_ks2_bez_nds:
            try:
                summa_ks2_bez_nds_rub = object.smeta.summa_ks2_bez_nds.split('.')[0]
                summa_ks2_bez_nds_kop = object.smeta.summa_ks2_bez_nds.split('.')[1]
            except:
                try:
                    summa_ks2_bez_nds_rub = object.smeta.summa_ks2_bez_nds.split(',')[0]
                    summa_ks2_bez_nds_kop = object.smeta.summa_ks2_bez_nds.split(',')[1]
                except:
                    pass

        doc.tables[4].rows[0].cells[1].text = doc.tables[4].rows[0].cells[1].text.replace('полеСуммабезндс', str(summa_ks2_bez_nds_rub))
        doc.tables[4].rows[0].cells[5].text = doc.tables[4].rows[0].cells[5].text.replace('кк', str(summa_ks2_bez_nds_kop))
        doc.tables[4].rows[2].cells[2].text = doc.tables[4].rows[2].cells[2].text.replace('полеСуммабезндс', str(summa_ks2_bez_nds_rub))
        doc.tables[4].rows[2].cells[5].text = doc.tables[4].rows[2].cells[5].text.replace('кк', str(summa_ks2_bez_nds_kop))
        doc.save('КС 11-КС 14 - {}.docx'.format(object.name_object))