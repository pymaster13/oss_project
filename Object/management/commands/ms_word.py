import datetime

from django.core.management.base import BaseCommand
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from Object.models import *

names_in_docx = {'полеНазваниеобъекта':'name_object', 'полеКодобъекта':'kod_object', \
'полеЗаявитель':'zayv', 'полеНомерпроекта':'Nomer_proekt', 'полеГод':str(datetime.datetime.today().year),\
'полеФИОсварщикп':'svarshik1.fio', 'полеФИОсварщикс':'svarshik2.fio', 'полеДиаметрп':'svarshik1.diametr',\
'полеДиаметрс':'svarshik2.diametr','полеКолвоп':'svarshik1.kolvo','полеКолвос':'svarshik2.kolvo'}


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

class Command(BaseCommand):
    help = 'test MS_WORD'

    def handle(self, *args, **options):
        object = Object.objects.first()
        doc = docx.Document('ITD_template.docx')

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        style_bold = doc.styles.add_style('Bold_style', WD_STYLE_TYPE.PARAGRAPH)
        style_bold.font.bold = True
        style_bold.font.size = Pt(11)

        style_11 = doc.styles.add_style('Normal_11', WD_STYLE_TYPE.PARAGRAPH)
        style_11.font.size = Pt(11)

        new_paragraph = doc.add_paragraph()
        new_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:

                """
                ГАЗОПРОВОДЫ
                """

                if 'Списокгазопроводов' in run.text:

                    paragraph.text = ''

                    paragraph.add_run("{}\n\n".format(object.name_gazoprovod)).bold = True

                    paragraph.add_run("{}\n".format(object.name_podzem_stal_gazoprovod)).bold = True
                    for truba in object.gazoprovod_podzem_stal.truba.all():
                        paragraph.add_run("- Труба Ø{}x{} - {} ({})\n".format(truba.diametr,
                                                                       truba.x,
                                                                       truba.dlina,
                                                                        truba.prim))
                    paragraph.add_run("  Установлено:\n")

                    paragraph.add_run("    - Неразъемное соединение ПЭØ{}/СТØ{} - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.neraz_soed.PE,
                        object.gazoprovod_podzem_stal.neraz_soed.ST,
                        object.gazoprovod_podzem_stal.neraz_soed.kolvo))

                    paragraph.add_run("    - Контрольная трубка - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.kontrolnaya_trubka))

                    paragraph.add_run("    - Отвод ст. 90 гр. - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.otvod_90))

                    paragraph.add_run("    - Опознавательный знак - {} шт.\n".format(
                        object.gazoprovod_podzem_stal.opoznavat_znak))

                    paragraph.add_run("\n{}\n".format(object.name_podzem_poliet_gazoprovod)).bold = True

                    try:
                        for truba in object.gazoprovod_podzem_poliet.truba.all():
                            paragraph.add_run("- Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                                truba.diametr, truba.x, truba.dlina))
                    except:
                        for truba in object.gazoprovod_podzem_poliet.truba.all():
                            paragraph.add_run("- Труба ПЭ100 SDR11 Ø{}х{} – {} м\n".format(
                            truba.diametr, truba.x, truba.dlina))


                    paragraph.add_run("  Установлено:\n")

                    for mufta in object.gazoprovod_podzem_poliet.mufta.all():
                        paragraph.add_run("    - Муфта электросварная ПЭ100 SDR11 Ø{}х{} – {} шт.\n".format(
                            mufta.diametr, mufta.x, mufta.dlina))

                    paragraph.add_run("    - Отвод ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.otvod.diametr,
                        object.gazoprovod_podzem_poliet.otvod.kolvo))

                    paragraph.add_run("    - Тройник электросварной ПЭ100 SDR11 Ø{}xØ{}xØ{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.troinik.diametr1,
                        object.gazoprovod_podzem_poliet.troinik.diametr2,
                        object.gazoprovod_podzem_poliet.troinik.diametr3,
                        object.gazoprovod_podzem_poliet.troinik.kolvo))

                    paragraph.add_run("    - Заглушка электросварная ПЭ100 SDR11 Ø{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.zaglushka.diametr,
                        object.gazoprovod_podzem_poliet.zaglushka.kolvo))

                    paragraph.add_run("    - Лента сигнальная “Осторожно! ГАЗ!”  – {} м\n".format(
                        object.gazoprovod_podzem_poliet.lenta))

                    paragraph.add_run("    - Кран шаровой для подземной установки ПЭ 100 ГАЗ SDR11 Ø{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.kran.diametr,
                        object.gazoprovod_podzem_poliet.kran.kolvo))

                    paragraph.add_run("    - Опознавательный знак – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.znak))

                    paragraph.add_run("    - Седелка электросварная Ø{}xØ{}xØ{} – {} шт.\n".format(
                        object.gazoprovod_podzem_poliet.sedelka.diametr1,
                        object.gazoprovod_podzem_poliet.sedelka.diametr2,
                        object.gazoprovod_podzem_poliet.sedelka.diametr3,
                        object.gazoprovod_podzem_poliet.sedelka.kolvo))

                    paragraph.add_run("\n{}\n".format(object.name_nadzem_stal_gazoprovod)).bold = True


                    for truba in object.gazoprovod_nadzem_stal.truba.all():
                        paragraph.add_run("- Труба стальная Ø{}х{} – {} м\n".format(
                            truba.diametr, truba.x, truba.dlina))

                    paragraph.add_run("  Установлено:\n")

                    paragraph.add_run("    - Изолирующие соединения Ø{} – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.izolir_soed.diametr,
                        object.gazoprovod_nadzem_stal.izolir_soed.kolvo))

                    paragraph.add_run("    - Кран стальной Ø{} – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.kran_stal.diametr,
                        object.gazoprovod_nadzem_stal.kran_stal.kolvo))

                    paragraph.add_run("    - Отвод Ø{} – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.otvod.diametr,
                        object.gazoprovod_nadzem_stal.otvod.kolvo))

                    paragraph.add_run("    - Цокольный ввод ПЭØ{}/СТØ{} - {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.cokolnyi_vvod.PE,
                        object.gazoprovod_nadzem_stal.cokolnyi_vvod.ST,
                        object.gazoprovod_nadzem_stal.cokolnyi_vvod.kolvo))

                    paragraph.add_run("    - Крепление газопровода к кирпичной стене – {} шт.\n".format(
                        object.gazoprovod_nadzem_stal.kreplenie))

                    paragraph.add_run("    - Стойка под газопровод Ø{} – {} м - {} шт.\n\n".format(
                        object.gazoprovod_nadzem_stal.stoika.diametr,
                        object.gazoprovod_nadzem_stal.stoika.dlina,
                        object.gazoprovod_nadzem_stal.stoika.kolvo))

                    paragraph.add_run("Итого протяженность газопровода – {} м\n".format(
                        object.itogo)).bold = True

                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.style = style_11

                """
                СЕРТИФИКАТЫ
                """

                if 'Списоксертификатов' in run.text:

                    count = 1
                    paragraph.text = ''
                    for cert in object.certificates.all():
                        paragraph.add_run('{}. {}\n'.format(count, cert.name))
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        count += 1

                """
                Страница 1
                """

                for name in names_in_docx:
                    if name != 'полеГод':
                        try:
                            if len(name.split('.')) > 0:
                                itog_field = object
                                for index in range(0,len(names_in_docx[name].split('.'))):
                                    splitted_field = names_in_docx[name].split('.')[index]
                                    itog_field = getattr(itog_field,splitted_field)
                                run.text = run.text.replace(name,itog_field)
                            else:
                                run.text = run.text.replace(name,getattr(object,names_in_docx[name]))
                        except:
                            pass
                    else:
                        run.text = run.text.replace(name,names_in_docx[name])

        """
        Страница 3
        """
        try:
            doc.tables[1].rows[0].cells[1].text = doc.tables[1].rows[0].cells[1].text.replace('полеНомерпроекта', object.kod_object)
            doc.tables[1].rows[3].cells[0].text = doc.tables[1].rows[3].cells[0].text.replace('полеНазваниеобъекта', object.name_object)
            doc.tables[1].rows[3].cells[0].text = doc.tables[1].rows[3].cells[0].text.replace('полеКодобъекта', object.kod_object)
            for paragraph in doc.tables[1].rows[3].cells[0].paragraphs:
                paragraph.style = style_bold
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

        if object.svarshik1:
            doc.tables[4].rows[3].cells[6].text = doc.tables[4].rows[3].cells[6].text.replace(
                'полеФИОсварщикп', object.svarshik1.fio)
            doc.tables[4].rows[4].cells[2].text = doc.tables[4].rows[4].cells[2].text.replace(
                'полеДиаметрп', object.svarshik1.diametr)
            doc.tables[4].rows[4].cells[3].text = doc.tables[4].rows[4].cells[3].text.replace(
                'полеКолвоп', object.svarshik1.kolvo)
            doc.tables[4].rows[4].cells[4].text = doc.tables[4].rows[4].cells[4].text.replace(
                'полеДатап', datetime.datetime.strftime(object.svarshik1.date_svarki,"%d.%m.%Y"))
        if object.svarshik2:
            doc.tables[4].rows[4].cells[5].text = doc.tables[4].rows[4].cells[5].text.replace(
                'полеФИОсварщикс', object.svarshik2.fio)
            doc.tables[4].rows[5].cells[1].text = doc.tables[4].rows[5].cells[1].text.replace(
                'полеДиаметрс', object.svarshik2.diametr)
            doc.tables[4].rows[5].cells[2].text = doc.tables[4].rows[5].cells[2].text.replace(
                'полеКолвос', object.svarshik2.kolvo)
            doc.tables[4].rows[5].cells[3].text = doc.tables[4].rows[5].cells[3].text.replace(
                'полеДатас', datetime.datetime.strftime(object.svarshik2.date_svarki,"%d.%m.%Y"))

        doc.tables[4].rows[7].cells[4].text = doc.tables[4].rows[7].cells[4].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph2 = doc.tables[4].rows[7].cells[4].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph3 = doc.tables[4].rows[7].cells[4].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[4].rows[18].cells[4].text = doc.tables[4].rows[18].cells[4].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph4 = doc.tables[4].rows[18].cells[4].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph5 = doc.tables[4].rows[18].cells[4].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[4].rows[21].cells[5].text = doc.tables[4].rows[21].cells[5].text.replace(
            'полеДатазамера1', datetime.datetime.strftime(object.Data_zamera1,"%d.%m.%Y"))
        doc.tables[4].rows[23].cells[5].text = doc.tables[4].rows[23].cells[5].text.replace(
            'полеДатазамера2', datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))

        doc.tables[5].rows[3].cells[3].text = doc.tables[5].rows[3].cells[3].text.replace('полеДавление1', object.prover_davl)
        doc.tables[5].rows[3].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[5].rows[4].cells[5].text = doc.tables[5].rows[4].cells[5].text.replace('полеДавление2', object.prover_davl)
        doc.tables[5].rows[4].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[5].rows[6].cells[3].text = doc.tables[5].rows[6].cells[3].text.replace('полеДавление3', object.prover_davl)
        doc.tables[5].rows[6].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[5].rows[12].cells[0].text = doc.tables[5].rows[12].cells[0].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph4 = doc.tables[5].rows[12].cells[0].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph5 = doc.tables[5].rows[12].cells[0].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[5].rows[16].cells[0].text = doc.tables[5].rows[16].cells[0].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
        doc.tables[5].rows[16].cells[0].text = doc.tables[5].rows[16].cells[0].text.replace(
            'полеДатасоставленияобъекта', datetime.datetime.strftime(object.Data_sost_project,"%d.%m.%Y"))
        doc.tables[5].rows[16].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.tables[5].rows[20].cells[4].text = doc.tables[5].rows[20].cells[4].text.replace(
            'полеДатаразбивки', datetime.datetime.strftime(object.Data_razbiv,"%d.%m.%Y"))
        doc.tables[5].rows[21].cells[4].text = doc.tables[5].rows[21].cells[4].text.replace(
            'полеДатазамера2', datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))


        doc.tables[5].rows[23].cells[0].text = doc.tables[5].rows[23].cells[0].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph6 = doc.tables[5].rows[23].cells[0].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph7 = doc.tables[5].rows[23].cells[0].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER


        doc.tables[6].rows[1].cells[6].text = doc.tables[6].rows[1].cells[6].text.replace(
            'полеДатапродувки', datetime.datetime.strftime(object.Data_produv,"%d.%m.%Y"))
        doc.tables[6].rows[1].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[6].rows[3].cells[0].text = doc.tables[6].rows[3].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
        doc.tables[6].rows[3].cells[0].text = doc.tables[6].rows[3].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)

        doc.tables[6].rows[11].cells[0].text = doc.tables[6].rows[11].cells[0].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        doc.tables[6].rows[11].cells[0].text = doc.tables[6].rows[11].cells[0].text.replace(
            'полеТехнадзорФИО', object.tehnadzor.person.fio)

        doc.tables[7].rows[1].cells[0].text = doc.tables[7].rows[1].cells[0].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
        doc.tables[7].rows[1].cells[4].text = doc.tables[7].rows[1].cells[4].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)


        doc.tables[8].rows[1].cells[0].text = doc.tables[8].rows[1].cells[0].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph8 = doc.tables[8].rows[1].cells[0].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph8.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph9 = doc.tables[8].rows[1].cells[0].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[9].rows[0].cells[5].text = doc.tables[9].rows[0].cells[5].text.replace(
            'полеДатазамера2', datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
        doc.tables[9].rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[10].rows[0].cells[3].text = doc.tables[10].rows[0].cells[3].text.replace('полеНомерпроекта', object.Nomer_proekt)
        doc.tables[10].rows[2].cells[0].text = doc.tables[10].rows[2].cells[0].text.replace('полеНазваниеобъекта', object.name_object)
        doc.tables[10].rows[2].cells[0].text = doc.tables[10].rows[2].cells[0].text.replace('полеКодобъекта', object.kod_object)

        flag_svarshik1 = False
        flag_svarshik2 = False
        if object.svarshik1:
            if object.prover_davl == '0.3':
                doc.tables[11].rows[1].cells[0].text = 'Низкого давления (полиэтилен)'
            elif object.prover_davl == '0.6':
                doc.tables[11].rows[1].cells[0].text = 'Среднего давления (полиэтилен)'
            else:
                doc.tables[11].rows[1].cells[0].text = '{} (полиэтилен)'.format(object.davlenie_name.capitalize())

            doc.tables[11].rows[1].cells[1].text = "{} МПа".format(object.prover_davl)
            doc.tables[11].rows[1].cells[2].text = "{} - {}".format(datetime.datetime.strftime(object.Data_zamera1,"%d.%m.%Y"),
                    datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
            doc.tables[11].rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.tables[11].rows[1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            flag_svarshik1 = True


        if object.svarshik2:
            if object.prover_davl == '0.3':
                doc.tables[11].rows[2].cells[0].text = 'Низкого давления (сталь)'
            elif object.prover_davl == '0.6':
                doc.tables[11].rows[2].cells[0].text = 'Среднего давления (сталь)'
            else:
                doc.tables[11].rows[2].cells[0].text = '{} (сталь)'.format(object.davlenie_name.capitalize())

            doc.tables[11].rows[2].cells[1].text = "{} МПа".format(object.prover_davl)
            doc.tables[11].rows[2].cells[2].text = "{} - {}".format(datetime.datetime.strftime(object.Data_zamera1,"%d.%m.%Y"),
                    datetime.datetime.strftime(object.Data_zamera2,"%d.%m.%Y"))
            doc.tables[11].rows[2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.tables[11].rows[2].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            flag_svarshik2 = True

        if flag_svarshik1:
            remove_row(doc.tables[11],doc.tables[11].rows[1])
        if flag_svarshik2:
            remove_row(doc.tables[11],doc.tables[11].rows[2])

        doc.tables[12].rows[7].cells[0].text = doc.tables[12].rows[7].cells[0].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph10 = doc.tables[12].rows[7].cells[0].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph10.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph11 = doc.tables[12].rows[7].cells[0].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph11.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[13].rows[0].cells[5].text = doc.tables[13].rows[0].cells[5].text.replace(
            'полеДатаразбивки1', datetime.datetime.strftime(object.Data_razbiv,"%d.%m.%Y"))
        doc.tables[13].rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[14].rows[0].cells[0].text = doc.tables[14].rows[0].cells[0].text.replace(
            'полеДатаразбивки2', datetime.datetime.strftime(object.Data_razbiv,"%d.%m.%Y"))
        doc.tables[14].rows[1].cells[1].text = doc.tables[14].rows[1].cells[1].text.replace(
            'полеНомеробъекта', object.Nomer_proekt)
        doc.tables[14].rows[1].cells[4].text = doc.tables[14].rows[1].cells[4].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
        doc.tables[14].rows[4].cells[0].text = doc.tables[14].rows[4].cells[0].text.replace(
            'полеНазваниеобъекта', object.name_object)
        doc.tables[14].rows[4].cells[0].text = doc.tables[14].rows[4].cells[0].text.replace(
            'полеКодобъекта', object.kod_object)
        doc.tables[14].rows[4].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if object.district.name == "Автозаводский район":
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[15].cells[0].text.replace('полеЗаказчикразбив',
            'Главный инженер Автозаводского производственного управления ПАО «Газпром газораспределение Нижний Новгород»')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________       {}'.format('Камбаратов А.В'))
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif object.district.name == "Сормовский район":
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[15].cells[0].text.replace('полеЗаказчикразбив',
            'Начальник СПУ ПАО «Газпром газораспределение Нижний Новгород»')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________       {}'.format('Метелев М.Л.'))
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif object.district.name == "Нижегородский район":
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[16].cells[0].text.replace('полеЗаказчикразбив',
            'Главный инженер Нагорного производственного управления ПАО «Газпром газораспределение Нижний Новгород»')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________       {}'.format('Котихин А.Н.'))
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        else:
            doc.tables[14].rows[15].cells[0].text = doc.tables[14].rows[15].cells[0].text.replace('полеЗаказчикразбив', '')
            paragraph12 = doc.tables[14].rows[16].cells[0].add_paragraph('_________________________                             ')
            paragraph12.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[15].rows[0].cells[0].text = doc.tables[15].rows[0].cells[0].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph13 = doc.tables[15].rows[0].cells[0].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph13.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph14 = doc.tables[15].rows[0].cells[0].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph14.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[16].rows[1].cells[0].text = doc.tables[16].rows[1].cells[0].text.replace('полеНазваниеобъекта', object.name_object)
        doc.tables[16].rows[1].cells[0].text = doc.tables[16].rows[1].cells[0].text.replace('полеКодобъекта', object.kod_object)

        doc.tables[16].rows[3].cells[8].text = doc.tables[16].rows[3].cells[8].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[16].rows[3].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[16].rows[11].cells[9].text = doc.tables[16].rows[11].cells[9].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        doc.tables[16].rows[11].cells[9].text = doc.tables[16].rows[11].cells[9].text.replace(
            'полеТехнадзорФИО', object.tehnadzor.person.fio)

        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'Работы выполнены по проектно-сметной документации', '2.  Работы выполнены по проектно-сметной документации')
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
        doc.tables[16].rows[17].cells[7].text = doc.tables[16].rows[17].cells[7].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(object.Data_sost_project,"%d.%m.%Y"))
        doc.tables[16].rows[18].cells[7].text = doc.tables[16].rows[18].cells[7].text.replace(
            'При выполнении работ применены', '3.  При выполнении работ применены')

        kran_mufta_string = ''
        for kran_mufta in object.dop_dann_sharovyi_kran.all():
            substr = 'ПЭ100 SDR11 кран Ø{}; Муфта ПЭ 100 ГАЗ SDR11 Ø{}; '.format(kran_mufta.kran.diametr, kran_mufta.mufta.diametr)
            kran_mufta_string += substr
        doc.tables[16].rows[19].cells[7].text = doc.tables[16].rows[19].cells[7].text.replace('полеКранмуфта,', kran_mufta_string)

        doc.tables[16].rows[20].cells[7].text = doc.tables[16].rows[20].cells[7].text.replace(
            'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

        doc.tables[16].rows[22].cells[7].text = doc.tables[16].rows[22].cells[7].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[16].rows[22].cells[7].text = doc.tables[16].rows[22].cells[7].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[16].rows[22].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        doc.tables[16].rows[28].cells[7].text = doc.tables[16].rows[28].cells[7].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph15 = doc.tables[16].rows[28].cells[7].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph15.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph16 = doc.tables[16].rows[28].cells[7].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph16.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[17].rows[1].cells[0].text = doc.tables[17].rows[1].cells[0].text.replace('полеНазваниеобъекта', object.name_object)
        doc.tables[17].rows[1].cells[0].text = doc.tables[17].rows[1].cells[0].text.replace('полеКодобъекта', object.kod_object)

        doc.tables[17].rows[3].cells[8].text = doc.tables[17].rows[3].cells[8].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[17].rows[3].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[17].rows[11].cells[9].text = doc.tables[17].rows[11].cells[9].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        doc.tables[17].rows[11].cells[9].text = doc.tables[17].rows[11].cells[9].text.replace(
            'полеТехнадзорФИО', object.tehnadzor.person.fio)

        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'Работы выполнены по проектно-сметной документации', '2.  Работы выполнены по проектно-сметной документации')
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеНомерпроекта', object.Nomer_proekt)
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
        doc.tables[17].rows[17].cells[7].text = doc.tables[17].rows[17].cells[7].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(object.Data_sost_project,"%d.%m.%Y"))
        doc.tables[17].rows[18].cells[7].text = doc.tables[17].rows[18].cells[7].text.replace(
            'При выполнении работ применены', '3.  При выполнении работ применены')

        doc.tables[17].rows[21].cells[7].text = doc.tables[17].rows[21].cells[7].text.replace(
            'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

        doc.tables[17].rows[23].cells[7].text = doc.tables[17].rows[23].cells[7].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[17].rows[23].cells[7].text = doc.tables[17].rows[23].cells[7].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[17].rows[23].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        doc.tables[17].rows[29].cells[7].text = doc.tables[17].rows[29].cells[7].text.replace(
            'полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph17 = doc.tables[17].rows[29].cells[7].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph17.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph18 = doc.tables[17].rows[29].cells[7].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph18.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.tables[18].rows[1].cells[0].text = doc.tables[18].rows[1].cells[0].text.replace('полеНазваниеобъекта', object.name_object)
        doc.tables[18].rows[1].cells[0].text = doc.tables[18].rows[1].cells[0].text.replace('полеКодобъекта', object.kod_object)

        doc.tables[18].rows[3].cells[1].text = doc.tables[18].rows[3].cells[1].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[18].rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[19].rows[5].cells[0].text = doc.tables[19].rows[5].cells[0].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        doc.tables[19].rows[5].cells[0].text = doc.tables[19].rows[5].cells[0].text.replace('полеТехнадзорФИО', object.tehnadzor.person.fio)


        doc.tables[20].rows[3].cells[0].text = doc.tables[20].rows[3].cells[0].text.replace(
            'полеЗащитныйфутляр', 'Ø{}х{} L={} м'.format(object.zashitnyi_futlyar.diametr,
                                            object.zashitnyi_futlyar.x, object.zashitnyi_futlyar.dlina))

        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'Работы выполнены по проектно-сметной документации', '2.  Работы выполнены по проектно-сметной документации')
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace('полеНомерпроекта', object.Nomer_proekt)
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеПроектнаяорганизация', object.proektnaya_org)
        doc.tables[20].rows[4].cells[0].text = doc.tables[20].rows[4].cells[0].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(object.Data_sost_project,"%d.%m.%Y"))
        doc.tables[20].rows[5].cells[0].text = doc.tables[20].rows[5].cells[0].text.replace(
            'При выполнении работ применены', '3.  При выполнении работ применены')

        trubi = ''
        for kran_mufta in object.dop_dann_gazopr_v_zashit.all():
            substr = 'Труба ПЭ100 ГАЗ SDR11 Ø{}х{}; труба ст.Ø{}х{}; '.format(
                    kran_mufta.truba1.diametr, kran_mufta.truba1.x,
                    kran_mufta.truba2.diametr, kran_mufta.truba2.x)
            trubi += substr
        trubi_endpoint = trubi[:-2] + '.'

        doc.tables[20].rows[6].cells[0].text = doc.tables[20].rows[6].cells[0].text.replace('полеТрубазащитныйфутляр', trubi_endpoint)
        doc.tables[20].rows[7].cells[0].text = doc.tables[20].rows[7].cells[0].text.replace(
            'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

        doc.tables[21].rows[0].cells[0].text = doc.tables[21].rows[0].cells[0].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[21].rows[0].cells[0].text = doc.tables[21].rows[0].cells[0].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[21].rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        doc.tables[21].rows[6].cells[0].text = doc.tables[21].rows[6].cells[0].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph19 = doc.tables[21].rows[6].cells[0].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph19.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph20 = doc.tables[21].rows[6].cells[0].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph20.alignment = WD_ALIGN_PARAGRAPH.CENTER


        doc.tables[22].rows[0].cells[18].text = doc.tables[22].rows[0].cells[18].text.replace('полеНазваниеобъекта', object.name_object)
        doc.tables[22].rows[0].cells[18].text = doc.tables[22].rows[0].cells[18].text.replace('полеКодобъекта', object.kod_object)

        doc.tables[22].rows[3].cells[0].text = doc.tables[22].rows[3].cells[0].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[22].rows[3].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[22].rows[10].cells[8].text = doc.tables[22].rows[10].cells[8].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        doc.tables[22].rows[10].cells[8].text = doc.tables[22].rows[10].cells[8].text.replace('полеТехнадзорФИО', object.tehnadzor.person.fio)


        doc.tables[22].rows[14].cells[3].text = doc.tables[22].rows[14].cells[3].text.replace(
            'полеФутлярнавыходе', 'Ø{}х{} L={}м – {} шт'.format(object.futlyar_na_vyhode.diametr,
                                            object.futlyar_na_vyhode.x, object.futlyar_na_vyhode.dlina, object.futlyar_na_vyhode.kolvo))

        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'Работы выполнены по проектно-сметной документации', '2.  Работы выполнены по проектно-сметной документации')
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace('полеНомерпроекта', object.Nomer_proekt)
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace('полеПроектнаяорганизация', object.proektnaya_org)
        doc.tables[22].rows[14].cells[18].text = doc.tables[22].rows[14].cells[18].text.replace(
            'полеДатасоставленияпроекта', datetime.datetime.strftime(object.Data_sost_project,"%d.%m.%Y"))
        doc.tables[22].rows[15].cells[12].text = doc.tables[22].rows[15].cells[12].text.replace(
            'При выполнении работ применены', '3.  При выполнении работ применены')

        trubi = ''
        for truba in object.dop_dann_futlyar_na_vyhode.all():
            substr = 'Труба Ø{}х{}; '.format(truba.diametr, truba.x)
            trubi += substr

        doc.tables[22].rows[16].cells[11].text = doc.tables[22].rows[16].cells[11].text.replace('полеТрубафутлярнавыходе', trubi)
        doc.tables[22].rows[17].cells[10].text = doc.tables[22].rows[17].cells[10].text.replace(
            'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

        doc.tables[22].rows[18].cells[9].text = doc.tables[22].rows[18].cells[9].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[22].rows[18].cells[9].text = doc.tables[22].rows[18].cells[9].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[22].rows[18].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        doc.tables[22].rows[21].cells[18].text = doc.tables[22].rows[21].cells[18].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph21 = doc.tables[22].rows[21].cells[18].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph22 = doc.tables[22].rows[21].cells[18].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER



        doc.tables[23].rows[1].cells[0].text = doc.tables[23].rows[1].cells[0].text.replace('полеНазваниеобъекта', object.name_object)
        doc.tables[23].rows[1].cells[0].text = doc.tables[23].rows[1].cells[0].text.replace('полеКодобъекта', object.kod_object)

        doc.tables[23].rows[3].cells[1].text = doc.tables[23].rows[3].cells[1].text.replace(
            'полеДатаукладки', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[23].rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.tables[24].rows[5].cells[0].text = doc.tables[24].rows[5].cells[0].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        doc.tables[24].rows[5].cells[0].text = doc.tables[24].rows[5].cells[0].text.replace('полеТехнадзорФИО', object.tehnadzor.person.fio)


        doc.tables[25].rows[3].cells[0].text = doc.tables[25].rows[3].cells[0].text.replace(
            'полеОпораподгазопровод', 'монтаж опоры Ø{}х{} L={} м над уровнем земли под надземный газопровод низкого давления в количестве {} шт'.format(
            object.opora.diametr, object.opora.x, object.opora.dlina, object.opora.kolvo))

        doc.tables[25].rows[4].cells[0].text = doc.tables[25].rows[4].cells[0].text.replace(
            'Работы выполнены по проектно-сметной документации', '2.  Работы выполнены по проектно-сметной документации')
        doc.tables[25].rows[5].cells[0].text = doc.tables[25].rows[5].cells[0].text.replace(
            'При выполнении работ применены', '3.  При выполнении работ применены')

        trubi = ''
        for truba in object.dop_dann_ob_ustanovke_opor.all():
            substr = 'Труба Ø{}х{}; '.format(truba.diametr, truba.x)
            trubi += substr

        doc.tables[25].rows[6].cells[0].text = doc.tables[25].rows[6].cells[0].text.replace('полеТрубаопораподгаз', trubi)
        doc.tables[25].rows[7].cells[0].text = doc.tables[25].rows[7].cells[0].text.replace(
            'При выполнении работ отсутствуют', '4.  При выполнении работ отсутствуют')

        doc.tables[25].rows[8].cells[0].text = doc.tables[25].rows[8].cells[0].text.replace(
            'полеДатаукладки1', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[25].rows[8].cells[0].text = doc.tables[25].rows[8].cells[0].text.replace(
            'полеДатаукладки2', datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
        doc.tables[25].rows[8].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        doc.tables[25].rows[14].cells[0].text = doc.tables[25].rows[14].cells[0].text.replace('полеТехнадзордолжность', object.tehnadzor.person.post)
        paragraph23 = doc.tables[25].rows[14].cells[0].add_paragraph('_________________________       {}'.format(object.tehnadzor.person.fio))
        paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph24 = doc.tables[25].rows[14].cells[0].add_paragraph('(должность производителя работ, фамилия, имя, отчество)')
        paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save('ИТД - {}.docx'.format(object.name_object))
