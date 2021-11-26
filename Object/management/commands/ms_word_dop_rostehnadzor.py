import datetime
from io import BytesIO
from zipfile import ZipFile

from django.core.management.base import BaseCommand
from django.http import HttpResponse
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

from Object.models import *


def rename_row_in_table(row, number):
    row.cells[0].text = number
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

names_in_docx = {'полеНазваниеобъекта':'name_object', \
                 'полеКодобъекта':'kod_object', \
                 'полеЗаявитель':'zayv', \
                 'полеНомерпроекта':'Nomer_proekt', \
                 'полеГод':str(datetime.datetime.today().year),\
                 'полеФИОсварщикп':'svarshik1.fio', \
                 'полеФИОсварщикс':'svarshik2.fio', \
                 'полеДиаметрп':'svarshik1.diametr',\
                 'полеДиаметрс':'svarshik2.diametr',\
                 'полеКолвоп':'svarshik1.kolvo',\
                 'полеКолвос':'svarshik2.kolvo'}


class Command(BaseCommand):
    help = 'test MS_WORD'

    def handle(self, *args, **options):
        object = Object.objects.get(name_object='2')
        zipObj = ZipFile('dop_ITD.zip', 'w')
        doc = docx.Document('Перечень лиц, участвующих в строительстве.docx')

        for index, table in enumerate(doc.tables):
            for index, row in enumerate(table.rows):
                for index, cell in enumerate(row.cells):
                    if 'полеНазваниеобъекта' in cell.text:
                        cell.text = cell.text.replace('полеНазваниеобъекта', object.name_object)

        doc.save('Перечень лиц, участв-х в строительстве.docx')

        doc2 = docx.Document('Акт нивелировки.docx')

        for index, paragraph in enumerate(doc2.paragraphs):
            for index, run in enumerate(paragraph.runs):
                if 'полеДатаукладки' in run.text:
                    if object.Data_ukl:
                        run.text = run.text.replace('полеДатаукладки',
                            datetime.datetime.strftime(object.Data_ukl,"%d.%m.%Y"))
                    else:
                        run.text = run.text.replace('полеДатаукладки', '')

                if 'полеНазваниеобъекта'in run.text:
                    if object.name_object:
                        run.text = run.text.replace('полеНазваниеобъекта',
                            object.name_object)
                    else:
                        run.text = run.text.replace('полеНазваниеобъекта', '')

                if 'полеПроектнаяорг'in run.text:
                    if object.proektnaya_org:
                        run.text = run.text.replace('полеПроектнаяорг',
                            object.proektnaya_org)
                    else:
                        run.text = run.text.replace('полеПроектнаяорг', '')

                if 'полеТехндолжность' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.post:
                                run.text = run.text.replace('полеТехндолжность',
                                    object.tehnadzor.person.post)
                            else:
                                run.text = run.text.replace('полеТехндолжность', '')
                        else:
                            run.text = run.text.replace('полеТехндолжность', '')
                    else:
                        run.text = run.text.replace('полеТехндолжность', '')
                else:
                    run.text = run.text.replace('полеТехндолжность', '')

                if 'полеТехнФИО' in run.text:
                    if object.tehnadzor:
                        if object.tehnadzor.person:
                            if object.tehnadzor.person.fio:
                                run.text = run.text.replace('полеТехнФИО',
                                    object.tehnadzor.person.fio)
                            else:
                                run.text = run.text.replace('полеТехнФИО', '')
                        else:
                            run.text = run.text.replace('полеТехнФИО', '')
                    else:
                        run.text = run.text.replace('полеТехнФИО', '')
                else:
                    run.text = run.text.replace('полеТехнФИО', '')


                if 'полеПК' in run.text:
                    if object.rostehnadzor:
                        run.text = run.text.replace('полеПК',
                            f'ПК{object.rostehnadzor.pk1}+{object.rostehnadzor.pk1_diam}-ПК{object.rostehnadzor.pk2}+{object.rostehnadzor.pk2_diam}')
                    else:
                        run.text = run.text.replace('полеПК', 'ПК_+_-ПК_+_')
                else:
                    run.text = run.text.replace('полеПК', 'ПК_+_-ПК_+_')

        doc2 = doc2.save('Акт нивелир-ки.docx')
        zipObj.write('Перечень лиц, участв-х в строительстве.docx')
        zipObj.write('Акт нивелир-ки.docx')
        zipObj.close()

        f = BytesIO()
        f.seek(0)
        with ZipFile('dop_ITD.zip', mode='r') as zf:
            f = zf.read('Перечень лиц, участв-х в строительстве.docx')
            f = zf.read('Акт нивелир-ки.docx')

        response = HttpResponse(f)
        response['Content-Type'] = 'application/x-zip-compressed'
        response['Content-Disposition'] = 'attachment; filename=dop_ITD.zip'

        return response


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
